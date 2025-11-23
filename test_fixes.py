import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from pathlib import Path

print("=" * 80)
print("TEST 1: Check individual IDML file with DAF mode")
print("=" * 80)

# First, check if we need to regenerate
perek1_file = Path(r'output\שס\ביצה\PEREK1-formatted.docx')
if perek1_file.exists():
    doc = Document(perek1_file)
    
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    print("\nFirst 20 paragraphs:")
    print("-" * 80)
    
    heading_count = {'Heading 3': 0, 'Heading 4': 0}
    
    for i, para in enumerate(doc.paragraphs[:20]):
        style = para.style.name if para.style else "Normal"
        text = para.text[:80] if para.text else ""
        print(f"[{i}] {style}: {text}")
        
        if style in heading_count:
            heading_count[style] += 1
    
    print(f"\nHeading 3 count in first 20: {heading_count['Heading 3']}")
    print(f"Heading 4 count in first 20: {heading_count['Heading 4']}")
    
    # Check if it starts with דף
    if doc.paragraphs:
        first_text = doc.paragraphs[0].text
        first_style = doc.paragraphs[0].style.name
        print(f"\nFirst paragraph: {first_style} - {first_text[:50]}")
        
        if 'דף' in first_text and first_style == 'Heading 3':
            print("✓ DAF mode appears to be working! (דף found in Heading 3)")
        else:
            print("✗ DAF mode NOT working (expected דף in Heading 3)")

print("\n" + "=" * 80)
print("TEST 2: Check combined file")
print("=" * 80)

combined_file = Path(r'output\שס\ביצה\ביצה-combined.docx')
if combined_file.exists():
    doc = Document(combined_file)
    
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    print("\nFirst 20 paragraphs:")
    print("-" * 80)
    
    heading_positions = {}
    
    for i, para in enumerate(doc.paragraphs[:20]):
        style = para.style.name if para.style else "Normal"
        text = para.text[:80] if para.text else ""
        print(f"[{i}] {style}: {text}")
        
        if style.startswith('Heading'):
            if style not in heading_positions:
                heading_positions[style] = []
            heading_positions[style].append((i, text))
    
    print("\nHeading positions:")
    for style, positions in sorted(heading_positions.items()):
        print(f"  {style}: {len(positions)} occurrence(s)")
        for pos, text in positions:
            print(f"    Position {pos}: {text[:50]}")
    
    # Check if headings repeat unnecessarily
    if heading_positions.get('Heading 1', []):
        if len(heading_positions['Heading 1']) > 1:
            print("\n✗ Heading 1 appears multiple times (should only appear once)")
        else:
            print("\n✓ Heading 1 appears only once")


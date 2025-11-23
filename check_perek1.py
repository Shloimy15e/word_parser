import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

# Check PEREK1 file
doc = Document(r'output\שס\ביצה\PEREK1-formatted.docx')

print("First 30 paragraphs of PEREK1-formatted.docx:")
print("=" * 80)

for i, para in enumerate(doc.paragraphs[:30]):
    style = para.style.name if para.style else "Normal"
    text = para.text[:100] if para.text else ""
    print(f"[{i}] {style}: {text}")

print("\n" + "=" * 80)
print(f"Total paragraphs: {len(doc.paragraphs)}")

# Count headings
heading_counts = {}
for para in doc.paragraphs:
    style = para.style.name if para.style else "Normal"
    if "Heading" in style:
        heading_counts[style] = heading_counts.get(style, 0) + 1

print("\nHeading counts:")
for style, count in sorted(heading_counts.items()):
    print(f"  {style}: {count}")


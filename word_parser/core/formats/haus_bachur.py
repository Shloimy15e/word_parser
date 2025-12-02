"""
Haus-bachur format handler.

This format handles documents with a specific structure:
- H1: From argument (book name)
- H2: Short line with font size 22
- H3: Bold line above a dotted line
- H4: Non-bold short sentence line with a box character before it
"""

import re
from typing import Dict, Any, List

from word_parser.core.formats._base import (
    DocumentFormat,
    FormatRegistry,
    Document,
    HeadingLevel,
    remove_page_markings,
)


@FormatRegistry.register
class HausBachurFormat(DocumentFormat):
    """
    Haus-bachur document format.
    
    Structure:
    - H1: Book name (from argument)
    - H2: Short line with font size 22
    - H3: Bold line above a dotted line separator
    - H4: Non-bold short sentence line with a box character (□, ☐, etc.) before it
    """

    @classmethod
    def get_format_name(cls) -> str:
        return "haus-bachur"

    @classmethod
    def get_priority(cls) -> int:
        return 15  # Higher priority than standard format

    @classmethod
    def get_description(cls) -> str:
        return "Haus-bachur format: H2 from size 22 lines, H3 from bold lines above dotted lines, H4 from box-prefixed lines"

    @classmethod
    def detect(cls, doc: Document, context: Dict[str, Any]) -> bool:
        """Detect if document matches haus-bachur format."""
        # Check if explicitly requested
        if context.get("format") == "haus-bachur" or context.get("mode") == "haus-bachur":
            return True
        
        # Auto-detect: look for patterns characteristic of this format
        # 1. Look for font size 22 lines (potential H2)
        # 2. Look for bold lines followed by dotted lines (potential H3)
        # 3. Look for box characters (potential H4)
        
        has_size_22 = False
        has_bold_before_dotted = False
        has_box_character = False
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Check for font size 22 (H2 candidate)
            if not has_size_22 and len(para.runs) > 0:
                has_size_22 = any(
                    run.style.font_size is not None and abs(run.style.font_size - 22.0) < 0.5
                    for run in para.runs
                ) and cls._is_short_line(text)
            
            # Check for bold line followed by dotted line (H3 candidate)
            if not has_bold_before_dotted and i < len(doc.paragraphs) - 1:
                is_bold = any(
                    run.style.bold == True
                    for run in para.runs
                    if run.style.bold is not None
                )
                if is_bold:
                    next_para = doc.paragraphs[i + 1]
                    next_text = next_para.text.strip()
                    if cls._is_dotted_line(next_text):
                        has_bold_before_dotted = True
            
            # Check for box character (H4 candidate)
            if not has_box_character:
                # Common box characters: □ (U+25A1), ☐ (U+2610), ■ (U+25A0), ▢ (U+25A2)
                box_chars = ['□', '☐', '■', '▢', '\u25A1', '\u2610', '\u25A0', '\u25A2']
                if any(text.startswith(box) for box in box_chars):
                    has_box_character = True
        
        # If we find at least 1 of these patterns, likely this format
        # Font size 22 is a strong indicator, so even finding just that is enough
        pattern_count = sum([has_size_22, has_bold_before_dotted, has_box_character])
        return pattern_count >= 1

    @classmethod
    def get_required_context(cls) -> List[str]:
        return []  # Book is optional - can be provided or empty

    @classmethod
    def get_optional_context(cls) -> Dict[str, Any]:
        return {}

    def process(self, doc: Document, context: Dict[str, Any]) -> Document:
        """Process document according to haus-bachur format rules."""
        book = context.get("book", "")
        input_path = context.get("input_path", "")
        
        # Remove page markings first
        doc = remove_page_markings(doc)
        
        # Set H1 from argument
        doc.set_headings(h1=book)
        
        # Process the document to identify H2, H3, and H4 headings
        # Pass input_path so we can check for images in source document
        self._apply_haus_bachur_headings(doc, input_path)
        
        return doc

    def _apply_haus_bachur_headings(self, doc: Document, input_path: str = None) -> None:
        """Apply haus-bachur heading detection rules."""
        print(f"Haus-bachur: Processing {len(doc.paragraphs)} paragraphs")
        
        # Pre-build image map and textbox map from source document (only once)
        image_map = {}
        textbox_map = {}
        if input_path:
            print("Haus-bachur: Building image and textbox maps from source document...")
            image_map = self._build_image_map(input_path)
            textbox_map = self._build_textbox_map(input_path)
            print(f"Haus-bachur: Found {sum(1 for v in image_map.values() if v)} paragraphs with images")
            print(f"Haus-bachur: Found {sum(1 for v in textbox_map.values() if v)} paragraphs with textboxes")
        
        h2_count = 0
        h3_count = 0
        h4_count = 0
        
        i = 0
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # H2: Short line with font size 22 (check this first - highest priority)
            if text and self._is_font_size_22(para) and self._is_short_line(text):
                para.heading_level = HeadingLevel.HEADING_2
                h2_count += 1
                print(f"Haus-bachur: Detected H2 (size 22): '{text[:50]}'")
            
            # H3: Line with textbox AND size 13
            elif text and self._has_textbox(para, textbox_map) and self._is_font_size_13(para):
                para.heading_level = HeadingLevel.HEADING_3
                h3_count += 1
                print(f"Haus-bachur: Detected H3 (textbox + size 13): '{text[:50]}'")
            
            # H4: Non-bold single-line sentence with textbox (but NOT size 13, to distinguish from H3)
            elif text and self._has_textbox(para, textbox_map) and self._is_single_line_sentence(text) and not self._is_bold(para) and not self._is_font_size_13(para):
                para.heading_level = HeadingLevel.HEADING_4
                h4_count += 1
                print(f"Haus-bachur: Detected H4 (textbox + non-bold sentence, not size 13): '{text[:50]}'")
            
            i += 1
        
        print(f"Haus-bachur: Summary - H2: {h2_count}, H3: {h3_count}, H4: {h4_count}")

    @staticmethod
    def _is_short_line(text: str) -> bool:
        """Check if text is a short line (suitable for H2)."""
        # Short line: typically less than 100 characters
        return len(text) < 100

    @staticmethod
    def _is_dotted_line(text: str) -> bool:
        """Check if text is a dotted/dashed line separator."""
        if not text:
            return False
        # Pattern: 3+ repeating dashes, dots, underscores, or similar characters
        # Also check for various dash types: -, –, —, em dash, etc.
        # Allow some spaces between separators
        pattern1 = re.match(r"^[-=_~\.·]{3,}$", text)
        pattern2 = re.match(r"^[-\u2013\u2014\u2015=_.·\s]{3,}$", text)  # Includes em/en dashes and spaces
        # Check if it's mostly separator characters (at least 70%)
        if not pattern1 and not pattern2:
            separator_chars = sum(1 for c in text if c in ['-', '–', '—', '=', '_', '.', '·', '~', ' '])
            if len(text) >= 3 and separator_chars >= len(text) * 0.7:
                return True
        return bool(pattern1 or pattern2)

    def _is_font_size_22(self, para) -> bool:
        """Check if paragraph has font size 22."""
        if not para.runs:
            return False
        
        # Check if at least one run has font size 22 (or close to it, within 0.5pt)
        # Some paragraphs may have mixed font sizes, so we check if any run is size 22
        has_size_22 = any(
            run.style.font_size is not None and abs(run.style.font_size - 22.0) < 0.5
            for run in para.runs
        )
        
        return has_size_22
    
    def _is_font_size_13(self, para) -> bool:
        """Check if paragraph contains size 13 font (at least one run is size 13)."""
        if not para.runs:
            return False
        
        # Check if at least one run has size 13 (using font.size.pt)
        has_size_13 = False
        for run in para.runs:
            try:
                if run.font.size:
                    size_pt = run.font.size.pt
                    if abs(size_pt - 13.0) < 0.5:
                        has_size_13 = True
                        break
            except:
                pass
        
        if has_size_13:
            return True
        
        # Check XML for font size if runs don't have explicit size
        try:
            para_xml = para._element.xml
            # sz values are in half-points, so 13pt = 26
            sz_matches = re.findall(r'w:sz[^=]*="(\d+)"', para_xml)
            if sz_matches:
                sizes = [int(s) / 2 for s in sz_matches]
                if any(abs(s - 13.0) < 0.5 for s in sizes):
                    return True
        except:
            pass
        
        # Check style name - use substring match for styles containing "כותרת משנה חדש"
        if para.style_name and 'כותרת משנה חדש' in para.style_name:
            return True
        
        return False
    
    def _has_textbox(self, para, textbox_map: Dict[str, bool]) -> bool:
        """Check if paragraph contains a textbox using cached textbox map."""
        para_text = para.text.strip()
        if para_text and para_text in textbox_map:
            return textbox_map[para_text]
        
        # Fallback: check directly if not in map
        try:
            para_xml = para._element.xml
            return 'txbxContent' in para_xml or ('textbox' in para_xml.lower() and 'shape' in para_xml.lower())
        except:
            return False
    
    def _has_bottom_border(self, para, border_map: Dict[str, bool]) -> bool:
        """Check if paragraph has a bottom border using cached border map."""
        para_text = para.text.strip()
        if para_text and para_text in border_map:
            return border_map[para_text]
        return False
    
    def _has_drawing_or_pict(self, para) -> bool:
        """Check if paragraph contains drawing or pict elements (for H4 detection)."""
        try:
            para_xml = para._element.xml
            return 'w:drawing' in para_xml or 'w:pict' in para_xml
        except:
            return False

    def _is_bold(self, para, input_path: str = None) -> bool:
        """Check if paragraph contains bold text (including style-based bold)."""
        if not para.runs:
            return False
        
        # Check if runs have explicit bold
        has_explicit_bold = any(
            run.style.bold == True
            for run in para.runs
            if run.style.bold is not None
        )
        
        if has_explicit_bold:
            return True
        
        # Check source document XML for bold (style-based bold)
        if input_path:
            try:
                from docx import Document as DocxDocument
                from pathlib import Path as PathLib
                docx_doc = DocxDocument(str(PathLib(input_path)))
                
                # Find corresponding paragraph
                para_text = para.text.strip()
                for src_para in docx_doc.paragraphs:
                    if src_para.text.strip() == para_text:
                        # Check XML for bold
                        para_xml = src_para._element.xml
                        if 'w:b' in para_xml or 'w:bCs' in para_xml:
                            return True
                        break
            except Exception:
                pass
        
        # Known styles with bold
        known_bold_styles = ['כותרת משנה חדש', 'כותרת משנה[']
        if para.style_name in known_bold_styles:
            return True
        
        return False

    def _has_box_character(self, text: str) -> bool:
        """Check if text starts with a box character."""
        if not text:
            return False
        # Extended list of box characters and similar symbols
        # Common box characters: □ (U+25A1), ☐ (U+2610), ■ (U+25A0), ▢ (U+25A2)
        # Also check for other box-like characters
        box_chars = [
            '□', '☐', '■', '▢', '▣', '▤', '▥', '▦', '▧', '▨', '▩',
            '\u25A1', '\u2610', '\u25A0', '\u25A2', '\u25A3', '\u25A4', '\u25A5',
            '\u25A6', '\u25A7', '\u25A8', '\u25A9', '\u25AA', '\u25AB', '\u25AC',
            '\u25AD', '\u25AE', '\u25AF', '\u25B0', '\u25B1', '\u25B2', '\u25B3',
            '☑', '☒',  # Checkbox variants
        ]
        # Check if text starts with any box character (after trimming)
        text_stripped = text.lstrip()
        if text_stripped != text:
            # Had leading whitespace, check if box is after whitespace
            return any(text_stripped.startswith(box) for box in box_chars)
        return any(text.startswith(box) for box in box_chars)

    def _is_single_line_sentence(self, text: str) -> bool:
        """Check if text is a single-line sentence (suitable for H3/H4 after image)."""
        if not text:
            return False
        # Single line: no newlines, reasonable length (less than 200 chars)
        # Should have some content
        return '\n' not in text and len(text) > 0 and len(text) < 200
    
    def _is_short_sentence(self, text: str) -> bool:
        """Check if text is a short sentence (suitable for H4)."""
        # Short sentence: typically less than 150 characters
        # Should have some content (not just the box character)
        cleaned = self._remove_box_character(text).strip()
        return len(cleaned) > 0 and len(text) < 150

    def _build_textbox_map(self, input_path: str) -> Dict[str, bool]:
        """Build a map of paragraph text to textbox status from source document."""
        textbox_map = {}
        try:
            from docx import Document as DocxDocument
            from pathlib import Path as PathLib
            
            docx_doc = DocxDocument(str(PathLib(input_path)))
            
            # Build map of all paragraphs with textboxes
            for src_para in docx_doc.paragraphs:
                para_text = src_para.text.strip()
                if para_text:
                    has_textbox = False
                    try:
                        # Check XML for textbox elements
                        para_xml = src_para._element.xml
                        # Look for textbox-related elements
                        if 'txbxContent' in para_xml or ('textbox' in para_xml.lower() and 'shape' in para_xml.lower()):
                            has_textbox = True
                    except Exception:
                        pass
                    
                    textbox_map[para_text] = has_textbox
        except Exception:
            pass
        
        return textbox_map
    
    def _build_border_map(self, input_path: str) -> Dict[str, bool]:
        """Build a map of paragraph text to bottom border status from source document."""
        border_map = {}
        try:
            from docx import Document as DocxDocument
            from pathlib import Path as PathLib
            import xml.etree.ElementTree as ET
            
            docx_doc = DocxDocument(str(PathLib(input_path)))
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # First, build a map of styles with bottom borders
            styles_with_borders = set()
            for style in docx_doc.styles:
                try:
                    if hasattr(style, '_element'):
                        pPr = style._element.find('.//w:pPr', ns)
                        if pPr is not None:
                            pBdr = pPr.find('w:pBdr', ns)
                            if pBdr is not None:
                                bottom = pBdr.find('w:bottom', ns)
                                if bottom is not None:
                                    val = bottom.get('w:val')
                                    if val and val not in ['nil', 'none', None]:
                                        styles_with_borders.add(style.name)
                except Exception:
                    pass
            
            # Build map of all paragraphs with bottom borders
            for src_para in docx_doc.paragraphs:
                para_text = src_para.text.strip()
                if para_text:
                    has_bottom_border = False
                    try:
                        # Check paragraph format for bottom border
                        pf = src_para.paragraph_format
                        if pf.border_bottom and pf.border_bottom.val is not None:
                            has_bottom_border = True
                        
                        # Check paragraph XML for pBdr element
                        p_elem = src_para._element
                        pPr = p_elem.find('w:pPr', ns)
                        if pPr is not None:
                            pBdr = pPr.find('w:pBdr', ns)
                            if pBdr is not None:
                                bottom = pBdr.find('w:bottom', ns)
                                if bottom is not None:
                                    val = bottom.get('w:val')
                                    if val and val not in ['nil', 'none', None]:
                                        has_bottom_border = True
                        
                        # Check if paragraph style has border
                        if src_para.style and src_para.style.name in styles_with_borders:
                            has_bottom_border = True
                    except Exception:
                        pass
                    
                    border_map[para_text] = has_bottom_border
        except Exception:
            pass
        
        return border_map
    
    def _build_image_map(self, input_path: str) -> Dict[str, bool]:
        """Build a map of paragraph text to image status from source document."""
        image_map = {}
        try:
            from docx import Document as DocxDocument
            from pathlib import Path as PathLib
            
            docx_doc = DocxDocument(str(PathLib(input_path)))
            
            # Build map of all paragraphs with images
            for src_para in docx_doc.paragraphs:
                para_text = src_para.text.strip()
                if para_text:
                    has_image = False
                    try:
                        # Check paragraph level
                        if src_para._element.xpath('.//w:drawing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            has_image = True
                        elif src_para._element.xpath('.//w:pict', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            has_image = True
                        else:
                            # Check runs
                            for run in src_para.runs:
                                if run._element.xpath('.//w:drawing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                                    has_image = True
                                    break
                                if run._element.xpath('.//w:pict', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                                    has_image = True
                                    break
                    except Exception:
                        pass
                    
                    image_map[para_text] = has_image
        except Exception:
            pass
        
        return image_map
    
    def _has_image_cached(self, para, image_map: Dict[str, bool]) -> bool:
        """Check if paragraph contains an image using cached image map."""
        # Method 1: Check metadata if image info was stored
        if para.metadata.get('has_image'):
            return True
        
        # Method 2: Check cached image map
        para_text = para.text.strip()
        if para_text and para_text in image_map:
            return image_map[para_text]
        
        # Method 3: Fallback heuristic - paragraph with empty runs and short text
        has_empty_run = any(not run.text.strip() for run in para.runs)
        has_text = para.text.strip()
        
        # If paragraph has both empty runs AND text AND is a short sentence, might have image
        if has_empty_run and has_text and len(para.runs) > 1 and len(has_text) < 200:
            return True
        
        return False
    
    def _remove_box_character(self, text: str) -> str:
        """Remove box character from the beginning of text."""
        box_chars = ['□', '☐', '■', '▢', '\u25A1', '\u2610', '\u25A0', '\u25A2']
        for box in box_chars:
            if text.startswith(box):
                # Remove box and any following whitespace
                return text[len(box):].lstrip()
        return text


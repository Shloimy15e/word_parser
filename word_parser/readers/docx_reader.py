"""
Reader for Microsoft Word .docx files.
"""

from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from word_parser.core.document import (
    Document,
    Paragraph,
    TextRun,
    RunStyle,
    HeadingLevel,
    Alignment,
    Footnote,
)
from word_parser.readers.base import InputReader, ReaderRegistry


@ReaderRegistry.register
class DocxReader(InputReader):
    """Reader for Microsoft Word .docx files (Open XML format)."""

    @classmethod
    def get_extensions(cls) -> List[str]:
        return [".docx"]

    @classmethod
    def supports_file(cls, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    @classmethod
    def get_priority(cls) -> int:
        return 100  # Highest priority for .docx files

    def read(self, file_path: Path) -> Document:
        """Read a .docx file and return a Document object."""
        source = DocxDocument(str(file_path))
        doc = Document()
        doc.metadata.source_file = str(file_path)

        # Extract footnotes first (before processing paragraphs)
        self._extract_footnotes(source, doc)

        # Map original footnote IDs to our footnote IDs
        footnote_id_map = {}
        for fn in doc.footnotes:
            if fn.original_id is not None:
                footnote_id_map[fn.original_id] = fn.id

        # Process paragraphs and link footnote references
        for src_para in source.paragraphs:
            para = self._convert_paragraph(src_para, source, footnote_id_map)
            doc.paragraphs.append(para)

        return doc

    def _extract_footnotes(self, source: DocxDocument, doc: Document) -> None:
        """Extract footnotes from the document."""
        try:
            # Access footnotes through relationships
            # The relationship type for footnotes is: http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes
            FOOTNOTES_REL_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
            
            footnotes_part = None
            try:
                # Try to get footnotes relationship
                footnotes_rel = source.part.rels.get_by_reltype(FOOTNOTES_REL_TYPE)
                if footnotes_rel:
                    footnotes_part = footnotes_rel.target_part
            except (AttributeError, KeyError):
                # Try alternative: iterate through all relationships
                try:
                    for rel in source.part.rels.values():
                        if 'footnote' in rel.reltype.lower():
                            footnotes_part = rel.target_part
                            break
                except Exception:
                    pass
            
            if footnotes_part is None:
                print("Debug: No footnotes_part found in document")
                return
            
            # Get footnotes XML element
            try:
                # Access the XML element directly
                footnotes_xml = footnotes_part.element
            except AttributeError:
                # Try accessing via part's blob
                try:
                    from docx.oxml import parse_xml
                    footnotes_xml = parse_xml(footnotes_part.blob)
                except Exception as e:
                    print(f"Debug: Could not access footnotes XML: {e}")
                    return
            
            if footnotes_xml is None:
                print("Debug: footnotes_xml is None")
                return
            
            # Use the correct namespace
            NS_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            
            # Find all footnote elements directly (not with .// which searches descendants)
            footnote_elems = footnotes_xml.findall(f'{NS_W}footnote')
            print(f"Debug: Found {len(footnote_elems)} footnote elements")
            
            footnote_id = 1  # Our internal footnote ID counter
            
            # Iterate through footnote elements
            for footnote_elem in footnote_elems:
                # Get the footnote ID from the w:id attribute
                original_id_attr = footnote_elem.get(f'{NS_W}id')
                if original_id_attr is None:
                    print("Debug: Footnote element missing id attribute")
                    continue
                
                try:
                    original_id = int(original_id_attr)
                except (ValueError, TypeError):
                    print(f"Debug: Could not parse footnote ID: {original_id_attr}")
                    continue
                
                # Skip separator and continuation separator footnotes
                footnote_type = footnote_elem.get(f'{NS_W}type')
                if footnote_type in ('separator', 'continuationSeparator'):
                    print(f"Debug: Skipping separator footnote {original_id}")
                    continue
                
                # Extract footnote paragraphs
                footnote = Footnote(id=footnote_id, original_id=original_id)
                
                # Get all paragraphs in the footnote
                para_elems = footnote_elem.findall(f'.//{NS_W}p')
                
                for para_elem in para_elems:
                    # Convert footnote paragraph
                    footnote_para = self._convert_footnote_paragraph(para_elem)
                    if footnote_para:
                        footnote.paragraphs.append(footnote_para)
                
                if footnote.paragraphs:  # Only add if footnote has content
                    doc.add_footnote(footnote)
                    footnote_id += 1
                else:
                    print(f"Debug: Footnote {original_id} has no content paragraphs")
                    
        except Exception as e:
            # If footnote extraction fails, continue without footnotes
            import traceback
            print(f"Warning: Could not extract footnotes: {e}")
            traceback.print_exc()
    
    def _convert_footnote_paragraph(self, para_elem) -> Optional[Paragraph]:
        """Convert a footnote paragraph XML element to our Paragraph model."""
        para = Paragraph()
        para.format.alignment = Alignment.RIGHT
        para.format.right_to_left = True
        
        # Extract text from runs
        run_elems = para_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        for run_elem in run_elems:
            # Get text from t elements
            text_elems = run_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            text = ''.join(elem.text or '' for elem in text_elems)
            if text:
                para.add_run(text)
        
        return para if para.runs else None
    
    def _convert_paragraph(self, src_para, source: DocxDocument = None, footnote_id_map: dict = None) -> Paragraph:
        """Convert a python-docx paragraph to our Paragraph model."""
        para = Paragraph()

        # Convert alignment
        if src_para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            para.format.alignment = Alignment.CENTER
        elif src_para.alignment == WD_ALIGN_PARAGRAPH.LEFT:
            para.format.alignment = Alignment.LEFT
        elif src_para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            para.format.alignment = Alignment.JUSTIFY
        else:
            para.format.alignment = Alignment.RIGHT

        # Copy paragraph format
        pf = src_para.paragraph_format
        if pf.left_indent is not None:
            para.format.left_indent = (
                pf.left_indent.pt if hasattr(pf.left_indent, "pt") else pf.left_indent
            )
        if pf.right_indent is not None:
            para.format.right_indent = (
                pf.right_indent.pt
                if hasattr(pf.right_indent, "pt")
                else pf.right_indent
            )
        if pf.first_line_indent is not None:
            para.format.first_line_indent = (
                pf.first_line_indent.pt
                if hasattr(pf.first_line_indent, "pt")
                else pf.first_line_indent
            )
        if pf.space_before is not None:
            para.format.space_before = (
                pf.space_before.pt
                if hasattr(pf.space_before, "pt")
                else pf.space_before
            )
        if pf.space_after is not None:
            para.format.space_after = (
                pf.space_after.pt if hasattr(pf.space_after, "pt") else pf.space_after
            )
        para.format.line_spacing = pf.line_spacing
        para.format.keep_together = pf.keep_together
        para.format.keep_with_next = pf.keep_with_next
        para.format.page_break_before = pf.page_break_before
        para.format.widow_control = pf.widow_control

        # Store style name
        if src_para.style:
            para.style_name = src_para.style.name
            # Detect heading level from style
            if para.style_name == "Heading 1":
                para.heading_level = HeadingLevel.HEADING_1
            elif para.style_name == "Heading 2":
                para.heading_level = HeadingLevel.HEADING_2
            elif para.style_name == "Heading 3":
                para.heading_level = HeadingLevel.HEADING_3
            elif para.style_name == "Heading 4":
                para.heading_level = HeadingLevel.HEADING_4
        
        # Detect if this is an actual numbered list item (not just "List Paragraph" style)
        # Check if paragraph has numbering properties
        try:
            # Access the underlying XML element to check for numbering
            p_element = src_para._element
            numPr = p_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            if numPr is not None:
                # This paragraph has actual numbering (numbered list item)
                para.metadata['is_numbered_list'] = True
        except Exception:
            # If we can't check, rely on style name
            pass

        # Convert runs and detect footnote references
        for src_run in src_para.runs:
            run = self._convert_run(src_run, source, src_para, footnote_id_map)
            para.runs.append(run)

        return para

    def _convert_run(self, src_run, source: DocxDocument = None, src_para = None, footnote_id_map: dict = None) -> TextRun:
        """Convert a python-docx run to our TextRun model."""
        # Extract color if present
        color_rgb = None
        if src_run.font.color and src_run.font.color.rgb:
            try:
                # python-docx RGBColor might not be an int, but str(rgb) returns 'RRGGBB'
                rgb_obj = src_run.font.color.rgb
                hex_color = str(rgb_obj)
                if len(hex_color) == 6:
                    r = int(hex_color[0:2], 16)
                    g = int(hex_color[2:4], 16)
                    b = int(hex_color[4:6], 16)
                    color_rgb = (r, g, b)
            except Exception:
                # Fallback or ignore if color extraction fails
                pass

        style = RunStyle(
            bold=src_run.font.bold,
            italic=src_run.font.italic,
            underline=src_run.font.underline,
            font_size=src_run.font.size.pt if src_run.font.size else None,
            font_name=src_run.font.name,
            color_rgb=color_rgb,
            all_caps=src_run.font.all_caps,
            small_caps=src_run.font.small_caps,
            strike=src_run.font.strike,
            superscript=src_run.font.superscript,
            subscript=src_run.font.subscript,
        )

        # Check for footnote reference
        footnote_id = None
        if source and src_para and footnote_id_map:
            try:
                # Check if this run contains a footnote reference
                # Access the XML element directly
                run_xml = src_run._element
                NS_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                footnote_refs = run_xml.findall(f'.//{NS_W}footnoteReference')
                if footnote_refs:
                    # Get the first footnote reference ID
                    ref_id_attr = footnote_refs[0].get(f'{NS_W}id')
                    if ref_id_attr:
                        try:
                            original_id = int(ref_id_attr)
                            # Map to our internal footnote ID
                            footnote_id = footnote_id_map.get(original_id)
                        except (ValueError, TypeError) as e:
                            print(f"Debug: Could not parse footnote reference ID: {ref_id_attr}, error: {e}")
            except Exception as e:
                # If footnote reference detection fails, continue without it
                print(f"Debug: Error detecting footnote reference: {e}")
                pass

        return TextRun(text=src_run.text, style=style, footnote_id=footnote_id)


#!/usr/bin/env python3
"""
Word Parser - Reformat Hebrew DOCX files to standardized schema.

This is the main CLI entry point. It uses the modular reader/writer/format architecture
to support multiple input formats, output formats, and document schemas.

To add new input formats:
    See word_parser/readers/base.py for the InputReader interface.

To add new output formats:
    See word_parser/writers/base.py for the OutputWriter interface.

To add new document formats (schemas):
    See word_parser/core/formats.py for the DocumentFormat interface.
"""

import argparse
import copy
import traceback
from pathlib import Path
from typing import Optional

# Import the modular components
from word_parser.core.document import Document
from word_parser.core.processing import (
    is_old_header,
    should_start_content,
    extract_heading4_info,
    extract_daf_headings,
    extract_year,
)
from word_parser.core.formats import FormatRegistry
from word_parser.readers import ReaderRegistry
from word_parser.writers import WriterRegistry
from word_parser.utils import get_processable_files, get_file_stem
from word_parser.core.document import Document, Paragraph, HeadingLevel


class DocumentProcessor:
    """
    Main document processing class.

    Handles reading documents, applying format-specific transformations, and writing output.
    """

    def __init__(
        self,
        output_format: str = "docx",
        document_format: Optional[str] = None,
        special_heading: bool = False,
        font_size_heading: bool = False,
        chunking_strategy: str = "paragraph",
    ):
        """
        Initialize processor.

        Args:
            output_format: Output format name ('docx', 'json', etc.)
            document_format: Document format/schema name ('standard', 'daf', 'siman', etc.)
                           If None, auto-detection will be used.
            special_heading: Enable special heading detection mode.
            font_size_heading: Enable font-size-based heading detection mode.
            chunking_strategy: Chunking strategy for JSON output.
        """
        self.output_format = output_format
        self.document_format_name = document_format
        self.special_heading = special_heading
        self.font_size_heading = font_size_heading
        self.chunking_strategy = chunking_strategy

        self.writer = WriterRegistry.get_writer(output_format)
        if not self.writer:
            available = ", ".join(WriterRegistry.get_supported_formats())
            raise ValueError(
                f"Unknown output format: {output_format}. "
                f"Available formats: {available}"
            )

        # Get document format handler if specified
        self.format_handler = None
        if document_format:
            self.format_handler = FormatRegistry.get_format(document_format)
            if not self.format_handler:
                available = ", ".join(fmt["name"] for fmt in FormatRegistry.list_formats())
                raise ValueError(
                    f"Unknown document format: {document_format}. "
                    f"Available formats: {available}"
                )

    def _apply_format(self, doc: Document, context: dict) -> Document:
        """
        Apply document format processing.

        If a format was specified, uses that. Otherwise, auto-detects.

        Args:
            doc: Input document
            context: Processing context (headings, filename, etc.)

        Returns:
            Processed document
        """
        handler = self.format_handler

        # Auto-detect if no format specified
        if handler is None:
            handler = FormatRegistry.detect_format(doc, context)
            if handler:
                print(f"  (auto-detected format: {handler.get_format_name()})", end=" ")

        # Apply format processing
        if handler:
            doc = handler.process(doc, context)

        return doc

    def process_file(
        self,
        input_path: Path,
        output_path: Path,
        book: Optional[str] = None,
        sefer: Optional[str] = None,
        parshah: Optional[str] = None,
        filename: Optional[str] = None,
        skip_parshah_prefix: bool = False,
        use_filename_for_h4: bool = False,
    ) -> None:
        """
        Process a single file.

        Args:
            input_path: Path to input file
            output_path: Path for output file
            book: Book title (H1) - optional for formatted format
            sefer: Sefer/tractate title (H2) - optional for formatted format
            parshah: Parshah name (H3) - optional for formatted format
            filename: Optional filename info (H4)
            skip_parshah_prefix: Don't add 'פרשת' prefix
        """
        # Get reader for input file
        reader = ReaderRegistry.get_reader_for_file(input_path)
        if not reader:
            available = ", ".join(ReaderRegistry.get_supported_extensions()) or "none"
            raise ValueError(
                f"No reader found for: {input_path}. "
                f"Supported extensions: {available}"
            )

        # Read document
        doc = reader.read(input_path)

        # Set headings (use empty string if None for compatibility)
        doc.set_headings(h1=book or "", h2=sefer, h3=parshah, h4=filename)

        # Build context for format processing
        context = {
            "book": book,
            "sefer": sefer,
            "parshah": parshah,
            "filename": filename,
            "input_path": str(input_path),
            "skip_parshah_prefix": skip_parshah_prefix,
            "use_filename_for_h4": use_filename_for_h4,
            "special_heading": self.special_heading,
            "font_size_heading": self.font_size_heading,
        }

        # Apply document format processing
        doc = self._apply_format(doc, context)

        # Write output
        self.writer.write(
            doc,
            output_path,
            skip_parshah_prefix=skip_parshah_prefix,
            chunking_strategy=self.chunking_strategy,
        )

    def process_file_daf_mode(
        self,
        input_path: Path,
        output_path: Path,
        book: str,
        daf_folder: str,
        filename: str,
    ) -> None:
        """
        Process a single file in daf mode.

        Args:
            input_path: Path to input file
            output_path: Path for output file
            book: Book title (H1)
            daf_folder: Folder name (H2)
            filename: Filename for H3/H4 extraction
        """
        # Get reader for input file
        reader = ReaderRegistry.get_reader_for_file(input_path)
        if not reader:
            raise ValueError(f"No reader found for: {input_path}")

        # Read document
        doc = reader.read(input_path)

        # Extract headings from filename
        heading3, heading4 = extract_daf_headings(filename)

        # Set headings
        doc.set_headings(h1=book, h2=daf_folder, h3=heading3, h4=heading4)

        # Build context for format processing
        context = {
            "book": book,
            "daf_folder": daf_folder,
            "filename": filename,
            "input_path": str(input_path),
            "daf_mode": True,
            "special_heading": self.special_heading,
            "font_size_heading": self.font_size_heading,
        }

        # Apply document format processing (usually DafFormat in daf mode)
        doc = self._apply_format(doc, context)

        # Write output
        self.writer.write(doc, output_path)

    def get_output_extension(self) -> str:
        """Get the file extension for the current output format."""
        return self.writer.get_extension()


def process_single_file(args, file_path: Path, out_dir: Path) -> None:
    """Process a single file directly."""
    output_format = "json" if args.json else "docx"
    document_format = getattr(args, "format", None)
    processor = DocumentProcessor(
        output_format=output_format,
        document_format=document_format,
        special_heading=args.special_heading,
        font_size_heading=args.font_size_heading,
        chunking_strategy=args.json_chunking,
    )

    # Extract info from filename
    filename_stem = file_path.stem
    title = filename_stem.replace("-formatted", "")

    # Determine headings (for formatted format, these can be None)
    sefer = args.sefer if args.sefer else (file_path.parent.name if document_format != "formatted" else None)
    parshah = args.parshah if args.parshah else (title if document_format != "formatted" else None)

    # Determine heading4: use clean filename if option is set, otherwise extract year
    if getattr(args, "use_filename_for_h4", False):
        heading4 = title
    else:
        year = extract_year(filename_stem)
        heading4_info = extract_heading4_info(filename_stem)
        heading4 = year or heading4_info or title

    # Create output path
    out_dir.mkdir(parents=True, exist_ok=True)
    ext = processor.get_output_extension()

    if args.json:
        json_dir = out_dir / "json"
        json_dir.mkdir(parents=True, exist_ok=True)
        out_name = f"{filename_stem}.json"
        out_path = json_dir / out_name
    else:
        out_name = f"{filename_stem.replace('-formatted', '')}-formatted{ext}"
        out_path = out_dir / out_name

    print(f"📄 Processing single file: {file_path.name}")
    if args.book:
        print(f"   Book (H1): {args.book}")
    if sefer:
        print(f"   Sefer (H2): {sefer}")
    if parshah:
        print(f"   Section (H3): {parshah}")
    if heading4 != title:
        print(f"   Subsection (H4): {heading4}")
    print()

    try:
        print(f"   {file_path.name} → {out_path.name} ...", end=" ")
        processor.process_file(
            file_path,
            out_path,
            args.book,
            sefer,
            parshah,
            heading4,
            args.skip_parshah_prefix,
            getattr(args, "use_filename_for_h4", False),
        )
        print("✓ done")
        print(f"\n✅ Output saved to: {out_path}")
    except Exception as e:
        print(f"⚠️ error: {e}")
        traceback.print_exc()


def process_folder_structure(args, docs_dir: Path, out_dir: Path) -> None:
    """Process documents using folder structure mode."""
    output_format = "json" if args.json else "docx"
    document_format = getattr(args, "format", None)
    processor = DocumentProcessor(
        output_format=output_format,
        document_format=document_format,
        special_heading=args.special_heading,
        font_size_heading=args.font_size_heading,
        chunking_strategy=args.json_chunking,
    )

    sefer = docs_dir.name
    subdirs = [d for d in docs_dir.iterdir() if d.is_dir()]

    # If no subdirectories but files exist, process files directly (for formats like perek-h3)
    if not subdirs:
        # Process all file types for all formats
        files = get_processable_files(docs_dir, all_types=True)
        if files:
            print(f"📚 Processing folder: {sefer} (no subdirectories, processing files directly)\n")
            # Create output directory
            if args.json:
                out_subdir = out_dir / "json" / sefer
            else:
                out_subdir = out_dir / sefer
            out_subdir.mkdir(parents=True, exist_ok=True)
            
            total_success = 0
            for i, path in enumerate(files, 1):
                try:
                    filename_stem = get_file_stem(path)
                    title = filename_stem.replace("-formatted", "")
                    
                    # Determine heading4: use clean filename if option is set, otherwise extract year
                    if getattr(args, "use_filename_for_h4", False):
                        heading4 = title
                    else:
                        year = extract_year(title)
                        heading4_info = extract_heading4_info(title)
                        heading4 = year or heading4_info or title
                    
                    ext = processor.get_output_extension()
                    if args.json:
                        out_name = f"{filename_stem}.json"
                    else:
                        out_name = f"{filename_stem.replace('-formatted', '')}-formatted{ext}"
                    out_path = out_subdir / out_name
                    
                    file_display_name = path.stem if path.suffix else path.name
                    print(
                        f"  [{i}/{len(files)}] {file_display_name} → {out_path.name} ...",
                        end=" ",
                    )
                    
                    processor.process_file(
                        path,
                        out_path,
                        args.book,
                        sefer,
                        None,  # No parshah for perek-h3 format
                        heading4,
                        args.skip_parshah_prefix,
                        getattr(args, "use_filename_for_h4", False),
                    )
                    print("✓ done")
                    total_success += 1
                except Exception as e:
                    print(f"⚠️ error: {e}")
                    traceback.print_exc()
            
            print(f"\n✅ All done. Successfully processed {total_success}/{len(files)} file(s).")
            return
        
        print(f"No subdirectories or files found in {docs_dir}")
        return

    print(f"📚 Processing folder structure: {sefer}\n")
    total_success = 0
    total_files = 0

    for subdir in subdirs:
        parshah = subdir.name

        # Create output subdirectory
        if args.json:
            out_subdir = out_dir / "json" / sefer / parshah
        else:
            out_subdir = out_dir / sefer / parshah
        out_subdir.mkdir(parents=True, exist_ok=True)

        if args.combine_parshah:
            print(f"📂 Combining {parshah} ...")
            try:
                combine_parshah_docs(
                    processor,
                    subdir,
                    out_subdir,
                    args.book,
                    sefer,
                    parshah,
                    args.skip_parshah_prefix,
                    getattr(args, "use_filename_for_h4", False),
                )
                print("  ✓ done")
                total_success += 1
            except Exception as e:
                print(f"  ⚠️ error: {e}")
                traceback.print_exc()
            continue

        # Process all file types for all formats
        files = get_processable_files(subdir, all_types=True)
        if not files:
            continue

        print(f"📂 {parshah} ({len(files)} file(s))")

        for i, path in enumerate(files, 1):
            try:
                filename_stem = get_file_stem(path)
                title = filename_stem.replace("-formatted", "")

                # Determine heading4: use clean filename if option is set, otherwise extract year
                if getattr(args, "use_filename_for_h4", False):
                    heading4 = title
                else:
                    year = extract_year(title)
                    heading4_info = extract_heading4_info(title)
                    heading4 = year or heading4_info or title

                ext = processor.get_output_extension()
                if args.json:
                    out_name = f"{filename_stem}.json"
                else:
                    out_name = (
                        f"{filename_stem.replace('-formatted', '')}-formatted{ext}"
                    )
                out_path = out_subdir / out_name

                file_display_name = path.stem if path.suffix else path.name
                print(
                    f"  [{i}/{len(files)}] {file_display_name} → {out_path.name} ...",
                    end=" ",
                )

                processor.process_file(
                    path,
                    out_path,
                    args.book,
                    sefer,
                    parshah,
                    heading4,
                    args.skip_parshah_prefix,
                    getattr(args, "use_filename_for_h4", False),
                )
                print("✓ done")
                total_success += 1
                total_files += 1
            except Exception as e:
                print(f"⚠️ error: {e}")
                total_files += 1
        print()

    print(f"✅ All done. Successfully processed {total_success}/{total_files} file(s).")


def process_daf_mode(args, docs_dir: Path, out_dir: Path) -> None:
    """Process documents in daf mode."""
    output_format = "json" if args.json else "docx"
    # In daf mode, default to 'daf' document format unless overridden
    document_format = getattr(args, "format", None) or "daf"
    processor = DocumentProcessor(
        output_format=output_format,
        document_format=document_format,
        special_heading=args.special_heading,
        font_size_heading=args.font_size_heading,
        chunking_strategy=args.json_chunking,
    )

    # Heading 1: book arg if provided, otherwise parent folder name
    book_name = args.book if args.book else docs_dir.name

    # Get all subdirectories
    folder_dirs = [d for d in docs_dir.iterdir() if d.is_dir()]

    # If no subdirectories but files exist, process files directly (for formats like perek-h3)
    if not folder_dirs:
        files = get_processable_files(docs_dir)
        if files:
            print(f"📚 Processing in daf mode (no subdirectories, processing files directly)")
            print(f"   Book (H1): {book_name}\n")
            
            # Create output directory
            if args.json:
                out_subdir = out_dir / "json" / docs_dir.name
            else:
                out_subdir = out_dir / docs_dir.name
            out_subdir.mkdir(parents=True, exist_ok=True)
            
            total_success = 0
            for i, path in enumerate(files, 1):
                try:
                    filename_stem = get_file_stem(path)
                    title = filename_stem.replace("-formatted", "")
                    
                    ext = processor.get_output_extension()
                    if args.json:
                        out_name = f"{filename_stem}.json"
                    else:
                        out_name = f"{filename_stem.replace('-formatted', '')}-formatted{ext}"
                    out_path = out_subdir / out_name
                    
                    file_display_name = path.stem if path.suffix else path.name
                    print(
                        f"  [{i}/{len(files)}] {file_display_name} → {out_path.name} ...",
                        end=" ",
                    )
                    
                    processor.process_file_daf_mode(
                        path, out_path, book_name, docs_dir.name, title
                    )
                    print("✓ done")
                    total_success += 1
                except Exception as e:
                    print(f"⚠️ error: {e}")
                    traceback.print_exc()
            
            print(f"\n✅ All done. Successfully processed {total_success}/{len(files)} file(s).")
            return
        
        print(f"No subdirectories or files found in {docs_dir}")
        return

    print(f"📚 Processing in daf mode")
    print(f"   Book (H1): {book_name}\n")
    total_success = 0
    total_files = 0

    for folder_dir in folder_dirs:
        folder_name = folder_dir.name

        # Create output subdirectory
        if args.json:
            out_subdir = out_dir / "json" / docs_dir.name / folder_name
        else:
            out_subdir = out_dir / docs_dir.name / folder_name
        out_subdir.mkdir(parents=True, exist_ok=True)

        files = get_processable_files(folder_dir)
        if not files:
            continue

        if args.combine_parshah:
            print(f"📂 Combining {folder_name} ...")
            try:
                combine_parshah_docs_daf_mode(
                    processor,
                    folder_dir,
                    out_subdir,
                    book_name,
                    folder_name,
                )
                print("  ✓ done")
                total_success += 1
            except Exception as e:
                print(f"  ⚠️ error: {e}")
                traceback.print_exc()
            continue

        print(f"📂 {folder_name} ({len(files)} file(s))")

        for i, path in enumerate(files, 1):
            try:
                filename_stem = get_file_stem(path)
                title = filename_stem.replace("-formatted", "")

                ext = processor.get_output_extension()
                if args.json:
                    out_name = f"{filename_stem}.json"
                else:
                    out_name = (
                        f"{filename_stem.replace('-formatted', '')}-formatted{ext}"
                    )
                out_path = out_subdir / out_name

                file_display_name = path.stem if path.suffix else path.name
                print(
                    f"  [{i}/{len(files)}] {file_display_name} → {out_path.name} ...",
                    end=" ",
                )

                processor.process_file_daf_mode(
                    path, out_path, book_name, folder_name, title
                )
                print("✓ done")
                total_success += 1
                total_files += 1
            except Exception as e:
                print(f"⚠️ error: {e}")
                total_files += 1
        print()

    print(f"✅ All done. Successfully processed {total_success}/{total_files} file(s).")


def combine_parshah_docs(
    processor: DocumentProcessor,
    subdir: Path,
    out_subdir: Path,
    book: str,
    sefer: str,
    parshah: str,
    skip_parshah_prefix: bool,
    use_filename_for_h4: bool = False,
) -> None:
    """
    Combine all documents in a folder into one file.

    Headings are only added:
    1. The first time they appear
    2. When that specific heading level changes

    Args:
        processor: DocumentProcessor instance
        subdir: Input directory containing files to combine
        out_subdir: Output directory for combined file
        book: Book title (H1)
        sefer: Sefer/tractate title (H2)
        parshah: Parshah name (H3)
        skip_parshah_prefix: Don't add 'פרשת' prefix
    """
    files = get_processable_files(subdir)
    if not files:
        return

    # Create combined document
    combined_doc = Document()

    # Add initial headings at the beginning (H1, H2, H3)
    # These are known and should always appear at the start
    h3_val = (
        parshah if skip_parshah_prefix else (f"פרשת {parshah}" if parshah else None)
    )

    if book:
        combined_doc.add_paragraph(book, heading_level=HeadingLevel.HEADING_1)
    if sefer:
        combined_doc.add_paragraph(sefer, heading_level=HeadingLevel.HEADING_2)
    if h3_val:
        combined_doc.add_paragraph(h3_val, heading_level=HeadingLevel.HEADING_3)

    # Track last seen heading values (set to initial values)
    last_h1 = book
    last_h2 = sefer
    last_h3 = h3_val
    last_h4 = None

    # Process each file
    for file_path in files:
        # Get reader for input file
        reader = ReaderRegistry.get_reader_for_file(file_path)
        if not reader:
            continue

        # Read document
        doc = reader.read(file_path)

        # Extract info from filename
        filename_stem = get_file_stem(file_path)
        title = filename_stem.replace("-formatted", "")
        
        # Determine heading4: use clean filename if option is set, otherwise extract year
        if use_filename_for_h4:
            heading4 = title
        else:
            year = extract_year(filename_stem)
            heading4_info = extract_heading4_info(filename_stem)
            heading4 = year or heading4_info or title

        # Set headings on the document
        doc.set_headings(h1=book, h2=sefer, h3=parshah, h4=heading4)

        # Build context for format processing
        context = {
            "book": book,
            "sefer": sefer,
            "parshah": parshah,
            "filename": heading4,
            "input_path": str(file_path),
            "skip_parshah_prefix": skip_parshah_prefix,
            "use_filename_for_h4": use_filename_for_h4,
            "special_heading": processor.special_heading,
            "font_size_heading": processor.font_size_heading,
        }

        # Apply document format processing
        doc = processor._apply_format(doc, context)

        # Determine heading values (with prefix handling for H3)
        h1_val = doc.heading1
        h2_val = doc.heading2
        h3_val = (
            doc.heading3
            if skip_parshah_prefix
            else (f"פרשת {doc.heading3}" if doc.heading3 else None)
        )
        h4_val = doc.heading4

        # Add headings only when they change
        if h1_val and h1_val != last_h1:
            para = combined_doc.add_paragraph(
                h1_val, heading_level=HeadingLevel.HEADING_1
            )
            last_h1 = h1_val

        if h2_val and h2_val != last_h2:
            para = combined_doc.add_paragraph(
                h2_val, heading_level=HeadingLevel.HEADING_2
            )
            last_h2 = h2_val

        if h3_val and h3_val != last_h3:
            para = combined_doc.add_paragraph(
                h3_val, heading_level=HeadingLevel.HEADING_3
            )
            last_h3 = h3_val

        if h4_val and h4_val != last_h4:
            para = combined_doc.add_paragraph(
                h4_val, heading_level=HeadingLevel.HEADING_4
            )
            last_h4 = h4_val

        # Add all body paragraphs from this document
        for para in doc.paragraphs:
            # Skip heading paragraphs (we handle headings separately above)
            if para.heading_level != HeadingLevel.NORMAL:
                continue

            # Create a deep copy of the paragraph
            new_para = copy.deepcopy(para)
            new_para.heading_level = HeadingLevel.NORMAL
            combined_doc.paragraphs.append(new_para)

    # Set headings on combined document to None so writer doesn't add them at start
    # (we've already added them as paragraphs when they changed)
    combined_doc.set_headings(h1=None, h2=None, h3=None, h4=None)

    # Determine output filename
    ext = processor.get_output_extension()
    if processor.output_format == "json":
        out_name = f"{parshah}.json"
    else:
        out_name = f"{parshah}-combined{ext}"
    out_path = out_subdir / out_name

    # Write combined document
    processor.writer.write(
        combined_doc,
        out_path,
        skip_parshah_prefix=skip_parshah_prefix,
        chunking_strategy=processor.chunking_strategy,
    )


def combine_parshah_docs_daf_mode(
    processor: DocumentProcessor,
    folder_dir: Path,
    out_subdir: Path,
    book: str,
    folder_name: str,
) -> None:
    """
    Combine all documents in a folder into one file (daf mode).

    Headings are only added:
    1. The first time they appear
    2. When that specific heading level changes

    Args:
        processor: DocumentProcessor instance
        folder_dir: Input directory containing files to combine
        out_subdir: Output directory for combined file
        book: Book title (H1)
        folder_name: Folder name (H2)
    """
    files = get_processable_files(folder_dir)
    if not files:
        return

    # Create combined document
    combined_doc = Document()

    # Add initial headings at the beginning (H1, H2)
    # These are known and should always appear at the start
    if book:
        combined_doc.add_paragraph(book, heading_level=HeadingLevel.HEADING_1)
    if folder_name:
        combined_doc.add_paragraph(folder_name, heading_level=HeadingLevel.HEADING_2)

    # Track last seen heading values (set to initial values)
    last_h1 = book
    last_h2 = folder_name
    last_h3 = None
    last_h4 = None

    # Process each file
    for file_path in files:
        # Get reader for input file
        reader = ReaderRegistry.get_reader_for_file(file_path)
        if not reader:
            continue

        # Read document
        doc = reader.read(file_path)

        # Extract headings from filename
        filename_stem = get_file_stem(file_path)
        title = filename_stem.replace("-formatted", "")
        heading3, heading4 = extract_daf_headings(title)

        # Set headings on the document
        doc.set_headings(h1=book, h2=folder_name, h3=heading3, h4=heading4)

        # Build context for format processing
        context = {
            "book": book,
            "daf_folder": folder_name,
            "filename": title,
            "input_path": str(file_path),
            "daf_mode": True,
            "special_heading": processor.special_heading,
            "font_size_heading": processor.font_size_heading,
        }

        # Apply document format processing
        doc = processor._apply_format(doc, context)

        # Get heading values
        h1_val = doc.heading1
        h2_val = doc.heading2
        h3_val = doc.heading3
        h4_val = doc.heading4

        # Add headings only when they change (H1 and H2 already added, so skip if same)
        if h1_val and h1_val != last_h1:
            para = combined_doc.add_paragraph(
                h1_val, heading_level=HeadingLevel.HEADING_1
            )
            last_h1 = h1_val

        if h2_val and h2_val != last_h2:
            para = combined_doc.add_paragraph(
                h2_val, heading_level=HeadingLevel.HEADING_2
            )
            last_h2 = h2_val

        if h3_val and h3_val != last_h3:
            para = combined_doc.add_paragraph(
                h3_val, heading_level=HeadingLevel.HEADING_3
            )
            last_h3 = h3_val

        if h4_val and h4_val != last_h4:
            para = combined_doc.add_paragraph(
                h4_val, heading_level=HeadingLevel.HEADING_4
            )
            last_h4 = h4_val

        # Add all body paragraphs from this document
        for para in doc.paragraphs:
            # Skip heading paragraphs (we handle headings separately above)
            if para.heading_level != HeadingLevel.NORMAL:
                continue

            # Create a deep copy of the paragraph
            new_para = copy.deepcopy(para)
            new_para.heading_level = HeadingLevel.NORMAL
            combined_doc.paragraphs.append(new_para)

    # Set headings on combined document to None so writer doesn't add them at start
    # (we've already added them as paragraphs when they changed)
    combined_doc.set_headings(h1=None, h2=None, h3=None, h4=None)

    # Determine output filename
    ext = processor.get_output_extension()
    if processor.output_format == "json":
        out_name = f"{folder_name}.json"
    else:
        out_name = f"{folder_name}-combined{ext}"
    out_path = out_subdir / out_name

    # Write combined document
    processor.writer.write(
        combined_doc,
        out_path,
        chunking_strategy=processor.chunking_strategy,
    )


def process_seif_footnotes(args, out_dir: Path) -> None:
    """
    Process seif-footnotes mode: merge content and footnotes files by matching seif markers.
    
    This mode:
    - Does NOT require --book, --sefer, --parshah, or any other arguments
    - Preserves ALL headings, metadata, and formatting exactly as-is from the content file
    - Only merges footnotes by matching seif markers (Hebrew letters)
    - Does NOT apply any format processing or transformations
    
    Args:
        args: Parsed command line arguments
        out_dir: Output directory
    """
    if not args.content_file:
        print("Error: --content-file is required for seif-footnotes mode")
        return
    if not args.footnotes_file:
        print("Error: --footnotes-file is required for seif-footnotes mode")
        return
    
    content_path = Path(args.content_file)
    footnotes_path = Path(args.footnotes_file)
    
    if not content_path.exists():
        print(f"Error: Content file not found: {content_path}")
        return
    if not footnotes_path.exists():
        print(f"Error: Footnotes file not found: {footnotes_path}")
        return
    
    # Determine output format
    output_format = "json" if args.json else "docx"
    
    # Determine output path
    out_dir.mkdir(parents=True, exist_ok=True)
    if args.json:
        json_dir = out_dir / "json"
        json_dir.mkdir(parents=True, exist_ok=True)
        output_name = f"{content_path.stem}-merged.json"
        output_path = json_dir / output_name
    else:
        output_name = f"{content_path.stem}-merged.docx"
        output_path = out_dir / output_name
    
    print(f"📄 Merging content and footnotes files (seif-footnotes mode)")
    print(f"   Content file: {content_path.name}")
    print(f"   Footnotes file: {footnotes_path.name}")
    print(f"   Output: {output_path.name}")
    print(f"   Note: All headings and formatting preserved exactly as-is")
    print()
    
    try:
        # Get the seif-footnotes writer
        writer = WriterRegistry.get_writer("seif-footnotes")
        if not writer:
            print("Error: seif-footnotes writer not found")
            return
        
        # Create a dummy document (will be ignored by the writer)
        dummy_doc = Document()
        
        # Write merged document
        # Note: We only pass output_format and chunking_strategy (for JSON)
        # We do NOT pass skip_parshah_prefix or any other format-related options
        # The writer preserves everything exactly as-is from the content file
        print(f"   Merging files...", end=" ")
        writer.write(
            dummy_doc,
            output_path,
            content_file=str(content_path),
            footnotes_file=str(footnotes_path),
            output_format=output_format,
            chunking_strategy=getattr(args, "json_chunking", "paragraph"),
        )
        print("✓ done")
        print(f"\n✅ Merged document saved to: {output_path}")
    except Exception as e:
        print(f"⚠️ error: {e}")
        traceback.print_exc()


def merge_docx_files(input_folder: Path, output_path: Path) -> None:
    """
    Merge all docx files in a folder (and subfolders) into one document.
    
    This is a simple merge that keeps all content exactly as-is.
    No processing, no heading extraction, no transformations.
    Just concatenates all documents in alphabetical order.
    Recursively searches subfolders and sub-subfolders.
    
    Args:
        input_folder: Folder containing docx files to merge
        output_path: Path for the merged output file
    """
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    # Get all docx files recursively, sorted alphabetically by full path
    files = sorted(input_folder.rglob("*.docx"))
    if not files:
        print(f"No .docx files found in {input_folder} (including subfolders)")
        return
    
    print(f"📂 Merging {len(files)} docx file(s) from {input_folder.name} (including subfolders)")
    print()
    
    # Create new document starting with the first file
    merged_doc = DocxDocument(str(files[0]))
    print(f"  [1/{len(files)}] {files[0].name} (base document)")
    
    # Append content from remaining files
    for i, file_path in enumerate(files[1:], 2):
        print(f"  [{i}/{len(files)}] {file_path.name}")
        
        # Read the source document
        source_doc = DocxDocument(str(file_path))
        
        # Add a page break before each new document
        merged_doc.add_page_break()
        
        # Copy all paragraphs from source to merged doc
        for para in source_doc.paragraphs:
            # Create new paragraph with same style
            new_para = merged_doc.add_paragraph()
            new_para.style = para.style
            new_para.alignment = para.alignment
            
            # Copy paragraph format properties
            if para.paragraph_format.first_line_indent:
                new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
            if para.paragraph_format.left_indent:
                new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
            if para.paragraph_format.right_indent:
                new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
            if para.paragraph_format.space_before:
                new_para.paragraph_format.space_before = para.paragraph_format.space_before
            if para.paragraph_format.space_after:
                new_para.paragraph_format.space_after = para.paragraph_format.space_after
            
            # Copy runs with their formatting
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                # Copy run formatting
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.color and run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
    
    # Save merged document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    merged_doc.save(str(output_path))
    
    print()
    print(f"✅ Merged document saved to: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description="Reformat Hebrew DOCX files to standardized schema.",
        epilog="""
Supported input formats: .docx, .doc (requires Word), .idml, DOS-encoded Hebrew
Supported output formats: docx, json

Special modes:
  --seif-footnotes: Merge content and footnotes files by matching seif markers
                    (Hebrew letters like א, ב, ג). Requires --content-file and --footnotes-file.

To add new formats, see the word_parser.readers and word_parser.writers packages.
        """,
    )
    parser.add_argument(
        "--book", help="Book title (Heading 1). Required unless using --daf mode."
    )
    parser.add_argument(
        "--sefer",
        help="Sefer/tractate title (Heading 2). If not provided, uses folder name.",
    )
    parser.add_argument(
        "--parshah",
        help="Parshah name (Heading 3). If not provided, uses subfolder names.",
    )
    parser.add_argument(
        "--skip-parshah-prefix",
        action="store_true",
        help="Skip adding 'פרשת' prefix to parshah name in Heading 3",
    )
    parser.add_argument(
        "--use-filename-for-h4",
        action="store_true",
        help="Use clean filename for H4 instead of extracting year (standard format only)",
    )
    parser.add_argument(
        "--json",
        action="store_true",
        help="Output as JSON structure instead of formatted Word documents",
    )
    parser.add_argument(
        "--docs",
        default="docs",
        help="Input folder containing files or subfolders (or single file for multi-parshah mode)",
    )
    parser.add_argument("--out", default="output", help="Output folder")
    parser.add_argument(
        "--multi-parshah",
        action="store_true",
        help="Process a single document containing multiple parshahs.",
    )
    parser.add_argument(
        "--combine-parshah",
        action="store_true",
        help="Combine all documents per folder into one file.",
    )
    parser.add_argument(
        "--daf",
        action="store_true",
        help="Daf mode: Parent folder → H1, Folder → H2, File name → H3/H4.",
    )
    parser.add_argument(
        "--special-heading",
        action="store_true",
        help="Special heading mode: H3 determined by preceding 'word.' or '– word –' or 'word – [word]' line",
    )
    parser.add_argument(
        "--font-size-heading",
        action="store_true",
        help="Font size heading mode (multi-parshah): H3 determined by size 14 standalone sentences",
    )
    parser.add_argument(
        "--list-formats",
        action="store_true",
        help="List all supported input, output, and document formats.",
    )
    parser.add_argument(
        "--format",
        help="Document format/schema (e.g., standard, daf, siman, multi-parshah, perek-h2). Auto-detected if not specified.",
    )

    parser.add_argument(
        "--json-chunking",
        choices=["paragraph", "h4", "h3", "chunk"],
        default="paragraph",
        help="Chunking strategy for JSON output: paragraph (default), h4, h3, or chunk (chunks within each H3 by asterisk markers).",
    )
    parser.add_argument(
        "--content-file",
        help="Content file for seif-footnotes mode (contains text with #(א) references)",
    )
    parser.add_argument(
        "--footnotes-file",
        help="Footnotes file for seif-footnotes mode (contains footnotes starting with 'א. ', 'ב. ', etc.)",
    )
    parser.add_argument(
        "--seif-footnotes",
        action="store_true",
        help="Merge content and footnotes files by matching seif markers (Hebrew letters). Requires --content-file and --footnotes-file.",
    )
    parser.add_argument(
        "--merge",
        nargs="?",
        const=True,
        default=False,
        metavar="FOLDER",
        help="Simple merge mode: combine all docx files in a folder into one file. No processing, keeps content exactly as-is. Optionally pass folder path directly.",
    )

    args = parser.parse_args()

    # Handle special heading mode shortcut
    # Only set format if not explicitly provided
    if args.special_heading and not args.format:
        args.format = "special-heading"

    # Handle --list-formats
    if args.list_formats:
        print("Supported input formats (file types):")
        for info in ReaderRegistry.list_readers():
            exts = (
                ", ".join(info["extensions"])
                if info["extensions"]
                else "(content-detected)"
            )
            print(f"  {info['name']}: {exts}")
        print("\nSupported output formats:")
        for info in WriterRegistry.list_writers():
            print(f"  {info['format']}: {info['extension']}")
        print("\nSupported document formats (schemas):")
        for info in FormatRegistry.list_formats():
            name = info["name"]
            # Get first line of description
            desc = (
                info["description"].strip().split("\n")[0]
                if info["description"]
                else ""
            )
            print(f"  {name}: {desc}")
        return

    # Handle --merge mode (simple docx merge, no processing)
    if args.merge:
        # Use path from --merge if provided, otherwise fall back to --docs
        if isinstance(args.merge, str):
            docs_path = Path(args.merge)
        else:
            docs_path = Path(args.docs)
        
        if not docs_path.exists():
            print(f"Error: Input path '{docs_path}' does not exist")
            return
        if not docs_path.is_dir():
            print(f"Error: --merge requires a folder path, got file: {docs_path}")
            return
        
        out_dir = Path(args.out)
        out_dir.mkdir(parents=True, exist_ok=True)
        output_path = out_dir / f"{docs_path.name}-merged.docx"
        
        merge_docx_files(docs_path, output_path)
        return

    # Validate --book is provided when not in daf mode or formatted format
    # Formatted format can extract headings from the document itself
    # Skip validation for seif-footnotes mode
    is_seif_footnotes_mode = args.seif_footnotes or (args.content_file and args.footnotes_file)
    if not is_seif_footnotes_mode and not args.daf and not args.book and getattr(args, "format", None) != "formatted" and getattr(args, "format", None) != "folder-filename":
        parser.error("--book is required unless using --daf mode or --format formatted or --format folder-filename")

    docs_path = Path(args.docs)
    out_dir = Path(args.out)

    # Seif-footnotes mode - merge content and footnotes files
    # This must come after out_dir is defined
    if is_seif_footnotes_mode:
        process_seif_footnotes(args, out_dir)
        return

    # Create output directory
    out_dir.mkdir(exist_ok=True)
    if args.json:
        (out_dir / "json").mkdir(exist_ok=True)

    # Multi-parshah mode
    if args.multi_parshah:
        print("⚠️ Multi-parshah mode not yet implemented in refactored version")
        return

    # Check input path
    if not docs_path.exists():
        print(f"Error: Input path '{docs_path}' does not exist")
        return

    # Single file mode - process one file directly
    if docs_path.is_file():
        process_single_file(args, docs_path, out_dir)
        return

    # Daf mode
    if args.daf:
        process_daf_mode(args, docs_path, out_dir)
        return

    # Folder structure mode (default when no sefer/parshah specified)
    if not args.sefer and not args.parshah:
        process_folder_structure(args, docs_path, out_dir)
        return

    # Check if format doesn't require parshah (like perek-h3)
    document_format = getattr(args, "format", None)
    formats_without_parshah = ["perek-h3", "perek-h2", "formatted"]
    format_doesnt_need_parshah = document_format in formats_without_parshah

    # Original single folder mode
    if not args.sefer or (not args.parshah and not format_doesnt_need_parshah):
        if format_doesnt_need_parshah:
            # For formats that don't need parshah, allow sefer-only
            pass  # Continue to process single folder
        else:
            print(
                "Error: Both --sefer and --parshah are required when not using folder structure mode"
            )
            return

    # Process single folder
    output_format = "json" if args.json else "docx"
    processor = DocumentProcessor(
        output_format=output_format,
        document_format=document_format,
        special_heading=args.special_heading,
        font_size_heading=args.font_size_heading,
        chunking_strategy=args.json_chunking,
    )

    files = get_processable_files(docs_path)
    if not files:
        print(f"No supported files found in {docs_path}")
        return

    print(f"📚 Processing {len(files)} file(s)...\n")

    success_count = 0
    for i, path in enumerate(files, 1):
        try:
            filename_stem = get_file_stem(path)
            title = filename_stem.replace("-formatted", "")

            # Determine heading4: use clean filename if option is set, otherwise extract year
            if getattr(args, "use_filename_for_h4", False):
                heading4 = title
            else:
                year = extract_year(filename_stem)
                heading4_info = extract_heading4_info(filename_stem)
                heading4 = year or heading4_info or title

            ext = processor.get_output_extension()
            if args.json:
                out_name = f"{filename_stem}.json"
                out_path = out_dir / "json" / out_name
            else:
                out_name = f"{filename_stem.replace('-formatted', '')}-formatted{ext}"
                out_path = out_dir / out_name

            file_display_name = path.stem if path.suffix else path.name
            print(
                f"[{i}/{len(files)}] Processing {file_display_name} → {out_path.name} ...",
                end=" ",
            )

            processor.process_file(
                path,
                out_path,
                args.book,
                args.sefer,
                args.parshah,
                heading4,
                args.skip_parshah_prefix,
                getattr(args, "use_filename_for_h4", False),
            )
            print("✓ done")
            success_count += 1
        except Exception as e:
            print(f"⚠️ error: {e}")
            traceback.print_exc()

    print(
        f"\n✅ All done. Successfully processed {success_count}/{len(files)} file(s)."
    )


if __name__ == "__main__":
    main()

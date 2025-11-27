"""
Writer for Microsoft Word .docx files.
"""

from pathlib import Path
from typing import Dict, Any

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from word_parser.core.document import Document, HeadingLevel, Alignment
from word_parser.core.processing import is_old_header, should_start_content
from word_parser.writers.base import OutputWriter, WriterRegistry


@WriterRegistry.register
class DocxWriter(OutputWriter):
    """Writer for Microsoft Word .docx files (Open XML format)."""

    @classmethod
    def get_format_name(cls) -> str:
        return "docx"

    @classmethod
    def get_extension(cls) -> str:
        return ".docx"

    @classmethod
    def get_default_options(cls) -> Dict[str, Any]:
        return {
            "skip_parshah_prefix": False,
            "filter_headers": True,
            "add_blank_lines": True,
        }

    def write(self, doc: Document, output_path: Path, **options) -> None:
        """
        Write document to a .docx file.

        Options:
            skip_parshah_prefix: Don't add 'פרשת' prefix to H3
            filter_headers: Skip old header paragraphs
            add_blank_lines: Add blank line after each paragraph
        """
        opts = {**self.get_default_options(), **options}

        # Create new document
        new_doc = DocxDocument()
        self._configure_styles(new_doc)

        # Add headings
        self._add_headings(new_doc, doc, opts.get("skip_parshah_prefix", False))

        # Process body paragraphs
        self._add_body_paragraphs(new_doc, doc, opts)

        # Add footnotes
        print(f"Debug: Document has {len(doc.footnotes)} footnotes")
        if doc.footnotes:
            print(f"Debug: Adding footnotes to output document...")
            self._add_footnotes(new_doc, doc)

        # Save
        output_path.parent.mkdir(parents=True, exist_ok=True)
        new_doc.save(str(output_path))

    def _configure_styles(self, docx_doc: DocxDocument) -> None:
        """Configure document styles for Hebrew text."""
        styles = docx_doc.styles

        def style_config(style_name, size, rgb, bold=True, space_after=6):
            try:
                s = styles[style_name]
                s.font.size = Pt(size)
                s.font.color.rgb = RGBColor(*rgb)
                s.font.bold = bold
                s.paragraph_format.space_after = Pt(space_after)
            except KeyError:
                pass

        style_config("Heading 1", 16, (0x2F, 0x54, 0x96), space_after=6)
        style_config("Heading 2", 13, (0x44, 0x72, 0xC4), space_after=4)
        style_config("Heading 3", 12, (0x1F, 0x37, 0x63), space_after=4)
        style_config("Heading 4", 11, (0x2F, 0x54, 0x96), space_after=4)

        try:
            normal = styles["Normal"]
            normal.font.size = Pt(12)
            normal.paragraph_format.space_after = Pt(0)
            normal.paragraph_format.line_spacing = 1.15
        except KeyError:
            pass

    def _format_parshah_heading(self, text: str, skip_prefix: bool) -> str:
        """
        Format parshah heading with or without prefix.

        Handles stripping existing prefix to avoid duplication or ensure removal.
        """
        if not text:
            return text

        # Remove existing prefix if present (to ensure clean slate)
        clean_text = text.strip()
        if clean_text.startswith("פרשת "):
            clean_text = clean_text[5:].strip()

        if skip_prefix:
            return clean_text
        else:
            return f"פרשת {clean_text}"

    def _add_headings(
        self, docx_doc: DocxDocument, doc: Document, skip_parshah_prefix: bool
    ) -> None:
        """Add document headings."""
        headings = []

        if doc.heading1:
            headings.append(("Heading 1", doc.heading1))
        if doc.heading2:
            headings.append(("Heading 2", doc.heading2))
        if doc.heading3:
            h3_text = self._format_parshah_heading(doc.heading3, skip_parshah_prefix)
            headings.append(("Heading 3", h3_text))
        if doc.heading4:
            headings.append(("Heading 4", doc.heading4))

        for level, text in headings:
            if text:
                p = docx_doc.add_paragraph(text, style=level)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.right_to_left = True

    def _add_body_paragraphs(
        self, docx_doc: DocxDocument, doc: Document, opts: Dict[str, Any]
    ) -> None:
        """Add body paragraphs with formatting."""
        filter_headers = opts.get("filter_headers", True)
        add_blank_lines = opts.get("add_blank_lines", True)
        is_multi_parshah = doc.metadata.extra.get("is_multi_parshah", False)

        in_header_section = filter_headers
        current_parshah = None  # Track current parshah for multi-parshah mode
        skip_parshah_prefix = opts.get("skip_parshah_prefix", False)

        for para in doc.paragraphs:
            txt = para.text.strip()

            # Handle heading paragraphs (for combined documents)
            if para.heading_level != HeadingLevel.NORMAL:
                # Map heading level to style name
                style_map = {
                    HeadingLevel.HEADING_1: "Heading 1",
                    HeadingLevel.HEADING_2: "Heading 2",
                    HeadingLevel.HEADING_3: "Heading 3",
                    HeadingLevel.HEADING_4: "Heading 4",
                }
                style_name = style_map.get(para.heading_level)
                if style_name and txt:
                    h = docx_doc.add_paragraph(txt, style=style_name)
                    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    h.paragraph_format.right_to_left = True
                continue  # Skip to next paragraph

            # Handle multi-parshah mode
            if is_multi_parshah:
                # Skip parshah marker lines (*, ה, etc.)
                if para.metadata.get("is_parshah_marker"):
                    continue

                # Check if this is a parshah boundary line (skip it, we'll add our own heading)
                if para.metadata.get("is_parshah_boundary"):
                    parshah_name = para.metadata.get("parshah_name", "")
                    if parshah_name and parshah_name != current_parshah:
                        current_parshah = parshah_name
                        # Add the parshah as a Heading 3
                        h3_text = self._format_parshah_heading(
                            parshah_name, skip_parshah_prefix
                        )
                        h3 = docx_doc.add_paragraph(h3_text, style="Heading 3")
                        h3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        h3.paragraph_format.right_to_left = True
                    continue  # Skip the original boundary paragraph

                # Check if parshah changed (for paragraphs that aren't boundaries)
                para_parshah = para.metadata.get("current_parshah")
                if para_parshah and para_parshah != current_parshah:
                    current_parshah = para_parshah
                    # This shouldn't normally happen, but handle it just in case

            # Header filtering logic
            if filter_headers and in_header_section:
                if txt and should_start_content(txt):
                    in_header_section = False
                elif txt and is_old_header(txt):
                    continue
                elif not txt:
                    continue
                else:
                    continue

            if filter_headers and txt and is_old_header(txt):
                continue

            # Skip paragraphs that are only markers (ה, *, ***, * * *)
            if txt in ("h", "q", "Y"):
                continue

            # Create new paragraph
            # If it's a list item, preserve the list style
            if para.is_list_item() and para.style_name:
                # Try to use the original list style if it exists
                try:
                    new_p = docx_doc.add_paragraph(style=para.style_name)
                except Exception:
                    # If style doesn't exist, create regular paragraph
                    new_p = docx_doc.add_paragraph()
            else:
                new_p = docx_doc.add_paragraph()

            # Set alignment
            if para.format.alignment == Alignment.CENTER:
                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif para.format.alignment == Alignment.LEFT:
                new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif para.format.alignment == Alignment.JUSTIFY:
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                new_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Set RTL
            new_p.paragraph_format.right_to_left = True

            # Copy paragraph formatting
            pf = new_p.paragraph_format
            if para.format.left_indent is not None:
                pf.left_indent = Pt(para.format.left_indent)
            if para.format.right_indent is not None:
                pf.right_indent = Pt(para.format.right_indent)
            if para.format.first_line_indent is not None:
                pf.first_line_indent = Pt(para.format.first_line_indent)
            if para.format.space_before is not None:
                pf.space_before = Pt(para.format.space_before)
            if para.format.space_after is not None:
                pf.space_after = Pt(para.format.space_after)
            if para.format.line_spacing is not None:
                pf.line_spacing = para.format.line_spacing
            if para.format.keep_together is not None:
                pf.keep_together = para.format.keep_together
            if para.format.keep_with_next is not None:
                pf.keep_with_next = para.format.keep_with_next
            if para.format.page_break_before is not None:
                pf.page_break_before = para.format.page_break_before
            if para.format.widow_control is not None:
                pf.widow_control = para.format.widow_control

            # Copy runs and add footnote references
            for run in para.runs:
                new_r = new_p.add_run(run.text)

                if run.style.bold is not None:
                    new_r.font.bold = run.style.bold
                if run.style.italic is not None:
                    new_r.font.italic = run.style.italic
                if run.style.underline is not None:
                    new_r.font.underline = run.style.underline
                if run.style.font_size is not None:
                    new_r.font.size = Pt(run.style.font_size)
                if run.style.font_name is not None:
                    new_r.font.name = run.style.font_name
                if run.style.color_rgb is not None:
                    new_r.font.color.rgb = RGBColor(*run.style.color_rgb)
                if run.style.all_caps is not None:
                    new_r.font.all_caps = run.style.all_caps
                if run.style.small_caps is not None:
                    new_r.font.small_caps = run.style.small_caps
                if run.style.strike is not None:
                    new_r.font.strike = run.style.strike
                if run.style.superscript is not None:
                    new_r.font.superscript = run.style.superscript
                if run.style.subscript is not None:
                    new_r.font.subscript = run.style.subscript
                
                # Add footnote reference if present
                if run.footnote_id is not None:
                    footnote = doc.get_footnote_by_id(run.footnote_id)
                    if footnote and footnote.original_id is not None:
                        # Add footnote reference using XML manipulation
                        self._add_footnote_reference(new_r, footnote.original_id)

            # Add blank line after non-empty paragraphs
            if add_blank_lines and txt:
                docx_doc.add_paragraph()
    
    def _add_footnote_reference(self, run, footnote_id: int) -> None:
        """Add a footnote reference to a run."""
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import qn
            
            # Create footnote reference element
            footnote_ref_xml = f'<w:footnoteReference xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:id="{footnote_id}"/>'
            footnote_ref = parse_xml(footnote_ref_xml)
            
            # Insert the footnote reference into the run
            run._element.append(footnote_ref)
        except Exception as e:
            # If footnote reference addition fails, continue without it
            print(f"Warning: Could not add footnote reference: {e}")
            pass
    
    def _add_footnotes(self, docx_doc: DocxDocument, doc: Document) -> None:
        """Add footnotes to the document."""
        try:
            from docx.oxml import parse_xml, OxmlElement
            from docx.oxml.ns import qn
            
            FOOTNOTES_REL_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
            
            # Get or create footnotes part
            footnotes_part = None
            footnotes_xml = None
            
            try:
                # Try to get existing footnotes part
                footnotes_rel = docx_doc.part.rels.get_by_reltype(FOOTNOTES_REL_TYPE)
                if footnotes_rel:
                    footnotes_part = footnotes_rel.target_part
                    # Get the element directly from the part
                    try:
                        footnotes_xml = footnotes_part._element
                    except AttributeError:
                        # If _element doesn't exist, parse from blob and set it
                        try:
                            footnotes_xml = parse_xml(footnotes_part.blob)
                            footnotes_part._element = footnotes_xml
                        except Exception:
                            # If we can't get the element, we'll create a new part
                            footnotes_part = None
                            footnotes_xml = None
            except (AttributeError, KeyError):
                pass
            
            # Create footnotes part if it doesn't exist
            if footnotes_part is None or footnotes_xml is None:
                print("Debug: Creating footnotes_part...")
                # Create footnotes XML structure with required separator footnotes
                # Word requires separator and continuationSeparator footnotes
                footnotes_xml = parse_xml(
                    '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
                    '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
                    '</w:footnotes>'
                )
                
                # Get the package
                package = docx_doc.part.package
                from docx.opc.constants import CONTENT_TYPE as CT
                from docx.opc.part import Part
                from docx.opc.packuri import PackURI
                
                # Create partname for footnotes
                partname = PackURI('/word/footnotes.xml')
                
                # Serialize the XML to bytes using lxml (python-docx uses lxml internally)
                # parse_xml returns an lxml element, so we must use lxml's tostring
                from lxml import etree as ET
                xml_bytes = ET.tostring(footnotes_xml, encoding='utf-8', xml_declaration=True, pretty_print=False)
                
                # Create new part using the package's method
                # The Part constructor will handle adding it to the package
                footnotes_part = Part(partname, CT.WML_FOOTNOTES, xml_bytes, package)
                
                # Add relationship
                docx_doc.part.relate_to(footnotes_part, FOOTNOTES_REL_TYPE)
                
                # Get the element from the part - try accessing via _element or parse blob
                try:
                    footnotes_xml = footnotes_part._element
                except AttributeError:
                    # Parse the blob to get the element
                    footnotes_xml = parse_xml(footnotes_part.blob)
                    # Set it as the part's element so modifications persist
                    footnotes_part._element = footnotes_xml
            
            print(f"Debug: Footnotes XML element: {footnotes_xml}, type: {type(footnotes_xml)}")
            
            # Add each footnote
            for footnote in doc.footnotes:
                if not footnote.paragraphs:
                    print(f"Debug: Skipping footnote {footnote.id} - no paragraphs")
                    continue
                
                footnote_id = footnote.original_id or footnote.id
                
                # Use OxmlElement for proper namespace handling
                footnote_elem = OxmlElement('w:footnote')
                footnote_elem.set(qn('w:id'), str(footnote_id))
                
                # Add paragraphs to footnote
                for para in footnote.paragraphs:
                    para_elem = OxmlElement('w:p')
                    
                    # Add runs to paragraph
                    for run in para.runs:
                        run_elem = OxmlElement('w:r')
                        
                        # Add text element
                        text_elem = OxmlElement('w:t')
                        # Set preserve space attribute
                        text_elem.set(qn('xml:space'), 'preserve')
                        # Set text content (lxml will handle XML escaping automatically)
                        if run.text:
                            text_elem.text = run.text
                        run_elem.append(text_elem)
                        para_elem.append(run_elem)
                    
                    footnote_elem.append(para_elem)
                
                # Check if footnote with this ID already exists (avoid duplicates)
                NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                existing_footnotes = footnotes_xml.findall(f'.//{{{NS_W}}}footnote')
                existing_ids = [fn.get(f'{{{NS_W}}}id') for fn in existing_footnotes if fn.get(f'{{{NS_W}}}id')]
                if str(footnote_id) not in existing_ids:
                    # Append footnote to footnotes XML (after separator footnotes if they exist)
                    footnotes_xml.append(footnote_elem)
                else:
                    print(f"Debug: Footnote {footnote_id} already exists, skipping")
            
            # Ensure the part's element is set to our modified XML and properly serialized
            if footnotes_part is not None:
                # Set the part's element to our modified XML
                footnotes_part._element = footnotes_xml
                
                # Serialize the XML properly using lxml (which python-docx uses internally)
                # parse_xml returns an lxml element, so we must use lxml's tostring
                from lxml import etree as ET
                xml_bytes = ET.tostring(
                    footnotes_xml,
                    encoding='utf-8',
                    xml_declaration=True,
                    pretty_print=False
                )
                object.__setattr__(footnotes_part, '_blob', xml_bytes)
            
            print(f"Debug: Finished adding {len(doc.footnotes)} footnotes to document")
                
        except Exception as e:
            # If footnote addition fails, continue without footnotes
            import traceback
            print(f"Warning: Could not add footnotes: {e}")
            traceback.print_exc()

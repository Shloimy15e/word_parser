import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

# Check the combined file
doc = Document(r'output\שס\ביצה\ביצה-combined.docx')

print("First 30 paragraphs of combined file:")
print("=" * 80)

for i, para in enumerate(doc.paragraphs[:30]):
    style = para.style.name if para.style else "Normal"
    text = para.text[:100] if para.text else ""
    print(f"[{i}] {style}: {text}")

print("\n" + "=" * 80)
print(f"Total paragraphs: {len(doc.paragraphs)}")


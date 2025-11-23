import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
from main import convert_to_docx
from docx import Document

# Test DAF mode on PEREK1.idml
idml_file = Path(r'backup_input\שס 2\ביצה\PEREK1.idml')

print("Testing DAF mode conversion...")
print("=" * 80)
print(f"Input: {idml_file}")

# Convert with DAF mode
output_path, needs_cleanup = convert_to_docx(idml_file, daf_mode=True)

print(f"Output: {output_path}")
print(f"Needs cleanup: {needs_cleanup}")

# Read the result
doc = Document(output_path)

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print("\nFirst 25 paragraphs:")
print("-" * 80)

daf_count = 0
amud_count = 0

for i, para in enumerate(doc.paragraphs[:25]):
    style = para.style.name if para.style else "Normal"
    text = para.text[:80] if para.text else ""
    print(f"[{i}] {style}: {text}")
    
    if style == 'Heading 3':
        daf_count += 1
    elif style == 'Heading 4':
        amud_count += 1

print("\n" + "=" * 80)
print(f"Heading 3 (דף) count in first 25: {daf_count}")
print(f"Heading 4 (עמוד) count in first 25: {amud_count}")

# Check first heading
if doc.paragraphs:
    first_text = doc.paragraphs[0].text
    first_style = doc.paragraphs[0].style.name
    print(f"\nFirst paragraph: {first_style} - {first_text}")
    
    if 'דף' in first_text and first_style == 'Heading 3':
        print("✓ DAF mode is working! (דף found in Heading 3)")
    else:
        print(f"✗ DAF mode NOT working")
        print(f"   Expected: Heading 3 with דף")
        print(f"   Got: {first_style} with {first_text}")

# Cleanup
if needs_cleanup:
    output_path.unlink()
    print("\n✓ Temp file cleaned up")


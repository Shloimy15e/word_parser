import re
import json
import argparse
import tempfile
import traceback
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

try:
    import win32com.client

    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False


# -------------------------------
# Helper: set up style formatting
# -------------------------------
def configure_styles(doc):
    styles = doc.styles

    def style_config(style_name, size, rgb, bold=True, space_after=6):
        try:
            s = styles[style_name]
            s.font.size = Pt(size)
            s.font.color.rgb = RGBColor(*rgb)
            s.font.bold = bold
            s.paragraph_format.space_after = Pt(space_after)
        except KeyError:
            # Style doesn't exist, skip it
            pass

    style_config("Heading 1", 16, (0x2F, 0x54, 0x96), space_after=6)
    style_config("Heading 2", 13, (0x44, 0x72, 0xC4), space_after=4)
    style_config("Heading 3", 12, (0x1F, 0x37, 0x63), space_after=4)
    style_config("Heading 4", 11, (0x2F, 0x54, 0x96), space_after=4)

    # Configure Normal style
    try:
        normal = styles["Normal"]
        normal.font.size = Pt(12)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.line_spacing = 1.15
    except KeyError:
        pass


# -------------------------------------------------
# Helper: decide which paragraphs are "old headers"
# -------------------------------------------------
HEADER_HINTS = [
    r"^דברות",
    r"^סדר",
    r"^פרשת",
    r"^שנת",
    r"^תש[\"׳]",
    r"^ס\"ג",
    r"^בעיר",
    r"^ב\"ה",
    r"^ליקוטי",
    r"^במסיבת",
    r"^מוצ\"ש",
    r"^מוצאי",
    r"^מוצש\"ק",
    r"^בבית.*התורה",
    r"^שבת",
    r"^פרשת.*שנת",
    r"^כ\"ק",
    r"לפ\"ק$",
    r"^יום.*פרשת.*שנת",
    r"^יום\s+[א-ת]['\"]",
]


def is_old_header(text):
    """
    Returns True if the paragraph looks like an old title/header line
    that should be skipped.
    """
    t = text.strip()
    if not t:
        return False  # Empty paragraphs should be preserved, not filtered

    # Single character paragraphs (like *) should be preserved
    if len(t) == 1:
        return False

    # Check against known header patterns
    if any(re.match(p, t) for p in HEADER_HINTS):
        return True

    # Skip short lines without punctuation (likely titles)
    # But NOT if it contains brackets [ ] which might be Torah text or single symbols
    if len(t) < 25 and not re.search(r"[.!?,\[\]\*]", t):
        return True

    return False


def should_start_content(text):
    """
    Returns True if this paragraph looks like substantive Torah content
    (long paragraph ≥60 chars OR contains Torah markers like brackets),
    signaling we're past the header section.
    """
    t = text.strip()
    # Torah content often has brackets for biblical quotes
    if "[" in t or "]" in t:
        return True
    # Or is a long paragraph
    return len(t) >= 60


# ----------------------------------------
# Helper: convert .doc to .docx using Word COM
# ----------------------------------------
def convert_doc_to_docx(doc_path):
    """
    Convert .doc file to .docx using Word COM automation.
    Returns path to temporary .docx file.
    """
    if not WORD_AVAILABLE:
        raise RuntimeError("pywin32 not installed. Cannot convert .doc files.")

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    try:
        # Open the .doc file
        doc = word.Documents.Open(str(doc_path.absolute()))

        # Create temp file for .docx
        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        temp_docx.close()

        # Save as .docx (format 16 = wdFormatXMLDocument)
        doc.SaveAs(temp_docx.name, FileFormat=16)
        doc.Close()

        return Path(temp_docx.name)
    finally:
        word.Quit()


def is_dos_encoded_file(file_path):
    """
    Check if a file is a DOS-encoded Hebrew text file (CP862).
    Returns True if it appears to be a DOS-encoded text file with Hebrew content.
    """
    if file_path.suffix:  # DOS files typically have no extension
        return False
    
    # Skip directories
    if not file_path.is_file():
        return False
    
    try:
        with open(file_path, "rb") as f:
            raw_data = f.read(2048)  # Read first 2KB for better detection
        
        # File must have some content
        if len(raw_data) == 0:
            return False
            
        # Try to decode as CP862 (Hebrew DOS)
        try:
            text = raw_data.decode('cp862', errors='strict')
            # Check if it contains Hebrew characters
            hebrew_chars = sum(1 for c in text if '\u0590' <= c <= '\u05FF')
            total_chars = len([c for c in text if c.isprintable() and not c.isspace()])
            
            # If more than 5% Hebrew characters (lowered threshold), likely a DOS Hebrew file
            # Also check that file is mostly text (not binary)
            if total_chars > 0 and hebrew_chars > total_chars * 0.05:
                return True
        except (UnicodeDecodeError, UnicodeError):
            # If strict decoding fails, try with errors='ignore'
            try:
                text = raw_data.decode('cp862', errors='ignore')
                hebrew_chars = sum(1 for c in text if '\u0590' <= c <= '\u05FF')
                # More lenient check with ignore errors
                if hebrew_chars > 10:  # At least 10 Hebrew characters
                    return True
            except:
                pass
            
        return False
    except Exception:
        return False


def sanitize_xml_text(text):
    """
    Remove characters that are not valid in XML.
    XML 1.0 valid characters: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD]
    """
    # Define valid XML character ranges
    def is_valid_xml_char(c):
        codepoint = ord(c)
        return (
            codepoint == 0x09 or  # Tab
            codepoint == 0x0A or  # Line feed
            codepoint == 0x0D or  # Carriage return
            (0x20 <= codepoint <= 0xD7FF) or
            (0xE000 <= codepoint <= 0xFFFD)
        )
    
    return ''.join(c for c in text if is_valid_xml_char(c))


def clean_dos_text(text):
    """
    Clean DOS text - remove ALL numbers, brackets, and formatting codes.
    Keep ONLY Hebrew text and basic punctuation.
    """
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line = line.strip()
        
        # Preserve empty lines
        if not line:
            cleaned_lines.append('')
            continue
        
        # Skip formatting lines starting with period
        if line.startswith('.'):
            continue
        
        # Must have Hebrew content
        if not any('\u0590' <= c <= '\u05FF' for c in line):
            continue
        
        temp = line
        
        # ============================================================
        # Remove ALL garbage - nuclear option
        # ============================================================
        
        # Remove >number< footnote markers
        temp = re.sub(r'>\d+<', '', temp)
        
        # Remove BNARF/OISAR/BSNF markers
        temp = re.sub(r'(BNARF|OISAR|BSNF)\s+[A-Z]\s+\d+[\*]?', '', temp)
        
        # Remove ALL brackets
        temp = re.sub(r'[<>]', '', temp)
        
        # Remove ALL numbers (integers and decimals)
        temp = re.sub(r'\d+\.?\d*', '', temp)
        
        # Remove asterisks
        temp = re.sub(r'\*', '', temp)
        
        # Remove multiple dashes
        temp = re.sub(r'[-–—]{2,}', '', temp)
        
        # Remove English letters (codes)
        temp = re.sub(r'[A-Za-z]+', '', temp)
        
        # Clean up spaces
        temp = re.sub(r'\s+', ' ', temp)
        temp = temp.strip()
        
        # Only keep if has Hebrew
        if temp and any('\u0590' <= c <= '\u05FF' for c in temp):
            cleaned_lines.append(temp)
    
    return '\n'.join(cleaned_lines)


def convert_dos_to_docx(dos_path):
    """
    Convert DOS-encoded Hebrew text file to .docx.
    Returns path to temporary .docx file.
    """
    # Read the DOS file
    with open(dos_path, "rb") as f:
        raw_data = f.read()
    
    # Decode from CP862 (Hebrew DOS encoding)
    text = raw_data.decode('cp862', errors='ignore')
    
    # Clean DOS formatting codes and garbage
    text = clean_dos_text(text)
    
    # Sanitize text to remove invalid XML characters
    text = sanitize_xml_text(text)
    
    # Create a new document
    doc = Document()
    
    # Split into paragraphs and add to document
    paragraphs = text.split('\n')
    for para_text in paragraphs:
        para_text = para_text.strip()
        if para_text:
            p = doc.add_paragraph(para_text)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.right_to_left = True
    
    # Save to temp file
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    temp_docx.close()
    doc.save(temp_docx.name)
    
    return Path(temp_docx.name)


def extract_text_from_idml(idml_path, daf_mode=False):
    """
    Extract text content from an IDML (InDesign Markup Language) file.
    IDML is a ZIP archive containing XML files.
    
    Args:
        idml_path: Path to the IDML file
        daf_mode: If True, returns list of (page_name, paragraph) tuples with page markers
                  If False, returns simple list of paragraph strings
    
    Returns:
        If daf_mode=True: list of (page_name, paragraph_text) tuples
        If daf_mode=False: list of paragraph strings
    """
    try:
        with zipfile.ZipFile(idml_path, 'r') as zip_file:
            # Find the main story (largest)
            story_files = [name for name in zip_file.namelist() if name.startswith('Stories/') and name.endswith('.xml')]
            largest_story = None
            max_size = 0
            main_story_id = None
            
            for story_file in story_files:
                info = zip_file.getinfo(story_file)
                if info.file_size > max_size:
                    max_size = info.file_size
                    largest_story = story_file
                    main_story_id = largest_story.split('_')[1].replace('.xml', '')
            
            if not largest_story:
                return [] if not daf_mode else []
            
            # Extract paragraphs from main story
            paragraphs = []
            with zip_file.open(largest_story) as f:
                tree = ET.parse(f)
                root = tree.getroot()
                
                for para_elem in root.iter():
                    tag_name = para_elem.tag.split('}')[-1] if '}' in para_elem.tag else para_elem.tag
                    
                    if tag_name == 'ParagraphStyleRange':
                        para_text_parts = []
                        for content_elem in para_elem.iter():
                            content_tag = content_elem.tag.split('}')[-1] if '}' in content_elem.tag else content_elem.tag
                            if content_tag == 'Content' and content_elem.text:
                                para_text_parts.append(content_elem.text)
                        
                        if para_text_parts:
                            para_text = ''.join(para_text_parts).strip()
                            para_text = para_text.replace('&apos;', "'").replace('&quot;', '"')
                            para_text = para_text.replace('\ufeff', '')
                            para_text = ' '.join(para_text.split())
                            
                            if para_text and para_text != "0" and len(para_text) > 1:
                                paragraphs.append(para_text)
            
            # If not in daf mode, return simple list
            if not daf_mode:
                return paragraphs
            
            # DAF MODE: Map paragraphs to pages
            frame_to_page = {}
            page_markers = {}
            
            # Parse spreads to find page mappings
            spread_files = [name for name in zip_file.namelist() if name.startswith('Spreads/') and name.endswith('.xml')]
            
            for spread_file in sorted(spread_files):
                with zip_file.open(spread_file) as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    
                    current_page = None
                    for elem in root.iter():
                        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                        
                        if tag == 'Page':
                            current_page = elem.attrib.get('Name', 'Unknown')
                        
                        elif tag == 'TextFrame' and current_page:
                            frame_id = elem.attrib.get('Self', '')
                            parent_story = elem.attrib.get('ParentStory', '')
                            
                            if frame_id:
                                frame_to_page[frame_id] = current_page
                                
                                # Check if this is a page marker (small story)
                                if parent_story != main_story_id and parent_story:
                                    try:
                                        story_file = f'Stories/Story_{parent_story}.xml'
                                        with zip_file.open(story_file) as sf:
                                            story_tree = ET.parse(sf)
                                            story_root = story_tree.getroot()
                                            texts = []
                                            for se in story_root.iter():
                                                st = se.tag.split('}')[-1] if '}' in se.tag else se.tag
                                                if st == 'Content' and se.text:
                                                    texts.append(se.text.strip())
                                            marker_text = ' '.join(texts).strip()
                                            # Store page marker if it looks like a reference
                                            if marker_text and ':' in marker_text and len(marker_text) < 50:
                                                if current_page not in page_markers:
                                                    page_markers[current_page] = marker_text
                                    except:
                                        pass
            
            # Get unique pages and sort by gematria value
            unique_pages = set(frame_to_page.values())
            pages_with_content = sorted(unique_pages, key=lambda p: hebrew_gematria_to_number(p))
            
            if not pages_with_content:
                # Fallback: return paragraphs without page info
                return [(None, p) for p in paragraphs]
            
            # Distribute paragraphs across pages (approximation)
            paras_per_page = len(paragraphs) // len(pages_with_content) if pages_with_content else len(paragraphs)
            
            result = []
            para_idx = 0
            
            for page_idx, page_name in enumerate(pages_with_content):
                # Add page marker if exists
                marker = page_markers.get(page_name, page_name)
                result.append((page_name, f"PAGE_MARKER:{marker}"))
                
                # Add paragraphs for this page
                end_idx = min(para_idx + paras_per_page, len(paragraphs))
                if page_idx == len(pages_with_content) - 1:
                    end_idx = len(paragraphs)  # Last page gets remaining
                
                for i in range(para_idx, end_idx):
                    result.append((page_name, paragraphs[i]))
                
                para_idx = end_idx
            
            return result
    
    except Exception as e:
        print(f"Warning: Error extracting text from IDML: {e}")
        import traceback
        traceback.print_exc()
        return [] if not daf_mode else []


def convert_idml_to_docx(idml_path, daf_mode=False):
    """
    Convert IDML file to .docx by extracting text content.
    
    Args:
        idml_path: Path to IDML file
        daf_mode: If True, uses page markers for Heading 3 (דף) and Heading 4 (עמוד)
    
    Returns:
        Path to temporary .docx file
    """
    # Extract text from IDML
    texts = extract_text_from_idml(idml_path, daf_mode=daf_mode)
    
    if not texts:
        raise ValueError(f"No text content found in IDML file: {idml_path}")
    
    # Create a new document
    doc = Document()
    
    if daf_mode:
        # Process with page markers - collect items first, then add in correct order
        current_daf = None
        current_amud = None
        pending_items = []  # List of (type, text, style) tuples
        
        def flush_pending():
            """Add pending items to document, reordering separators before headings"""
            if not pending_items:
                return
            
            # Find separators that come right after headings
            i = 0
            while i < len(pending_items):
                item_type, item_text, _ = pending_items[i]
                
                # If this is a separator and previous items are headings, move separator before them
                if item_type == 'separator':
                    # Look back to find where headings start
                    heading_start = i
                    for j in range(i - 1, -1, -1):
                        if pending_items[j][0] in ['heading3', 'heading4']:
                            heading_start = j
                        else:
                            break
                    
                    # If we found headings right before, move separator before them
                    if heading_start < i:
                        separator = pending_items.pop(i)
                        pending_items.insert(heading_start, separator)
                        i = heading_start + 1
                        continue
                
                i += 1
            
            # Now add all items to document
            for item_type, item_text, _ in pending_items:
                if item_type == 'heading3':
                    p = doc.add_paragraph(item_text, style="Heading 3")
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.right_to_left = True
                elif item_type == 'heading4':
                    p = doc.add_paragraph(item_text, style="Heading 4")
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.right_to_left = True
                elif item_type == 'separator':
                    p = doc.add_paragraph(item_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:  # regular paragraph
                    p = doc.add_paragraph(item_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.right_to_left = True
            
            pending_items.clear()
        
        for page_name, text in texts:
            if text.startswith("PAGE_MARKER:"):
                # Parse the marker to extract daf/amud info
                marker = text.replace("PAGE_MARKER:", "").strip()
                
                # Detect amud from punctuation: "." = amud 1, ":" = amud 2
                if '.' in marker:
                    amud = "א"
                elif ':' in marker:
                    amud = "ב"
                else:
                    amud = "א"  # Default to amud 1
                
                # Add heading for daf if changed
                if page_name and page_name != current_daf:
                    pending_items.append(('heading3', f"דף {page_name}", "Heading 3"))
                    current_daf = page_name
                    current_amud = None  # Reset amud when daf changes
                
                # Add heading for amud if changed
                if amud != current_amud:
                    pending_items.append(('heading4', f"עמוד {amud}", "Heading 4"))
                    current_amud = amud
            
            elif text.strip():
                # Skip "פרק" (chapter) headings in DAF mode
                if text.strip().startswith("פרק"):
                    continue
                
                # Check if this is a separator line (only asterisks and spaces)
                cleaned_text = text.strip().replace(' ', '').replace('\u00a0', '')
                is_separator = cleaned_text and all(c == '*' for c in cleaned_text)
                
                if is_separator:
                    pending_items.append(('separator', text, None))
                else:
                    pending_items.append(('paragraph', text, None))
        
        # Flush any remaining items
        flush_pending()
    else:
        # Simple mode: just add all text
        for text in texts:
            if text.strip():
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.right_to_left = True
    
    # Save to temp file
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    temp_docx.close()
    doc.save(temp_docx.name)
    
    return Path(temp_docx.name)


def get_processable_files(directory):
    """
    Get files to process from a directory.
    Priority order: .docx > .doc > .idml > DOS-encoded (no extension)
    Returns only files of ONE type (the highest priority type found).
    """
    files_by_type = {
        'docx': list(directory.glob("*.docx")),
        'doc': list(directory.glob("*.doc")),
        'idml': list(directory.glob("*.idml")),
        'dos': []
    }
    
    # Find DOS-encoded files (no extension)
    for file in directory.iterdir():
        if file.is_file() and not file.suffix and is_dos_encoded_file(file):
            files_by_type['dos'].append(file)
    
    # Return files in priority order
    for file_type in ['docx', 'doc', 'idml', 'dos']:
        if files_by_type[file_type]:
            return files_by_type[file_type]
    
    return []


def convert_to_docx(file_path, daf_mode=False):
    """
    Convert any supported file format to .docx.
    Returns tuple: (path_to_docx, needs_cleanup)
    - path_to_docx: Path object to the .docx file
    - needs_cleanup: Boolean indicating if the file is temporary and should be deleted
    """
    suffix = file_path.suffix.lower()
    
    if suffix == '.docx':
        # Already a .docx file, no conversion needed
        return file_path, False
    
    elif suffix == '.doc':
        # Convert .doc to .docx
        temp_docx = convert_doc_to_docx(file_path)
        return temp_docx, True
    
    elif suffix == '.idml':
        # Convert .idml to .docx
        temp_docx = convert_idml_to_docx(file_path, daf_mode=daf_mode)
        return temp_docx, True
    
    elif not suffix and is_dos_encoded_file(file_path):
        # Convert DOS-encoded file to .docx
        temp_docx = convert_dos_to_docx(file_path)
        return temp_docx, True
    
    else:
        raise ValueError(f"Unsupported file type: {file_path}")


# ----------------------------------------
# Core conversion: one docx → formatted docx
# ----------------------------------------
def hebrew_gematria_to_number(hebrew_str):
    """
    Convert Hebrew gematria notation to a number.
    Examples: א → 1, ב → 2, י → 10, יא → 11, טו → 15, etc.
    """
    if not hebrew_str:
        return 0
    
    # Hebrew letter values
    values = {
        'א': 1, 'ב': 2, 'ג': 3, 'ד': 4, 'ה': 5, 'ו': 6, 'ז': 7, 'ח': 8, 'ט': 9,
        'י': 10, 'כ': 20, 'ך': 20, 'ל': 30, 'מ': 40, 'ם': 40, 'נ': 50, 'ן': 50,
        'ס': 60, 'ע': 70, 'פ': 80, 'ף': 80, 'צ': 90, 'ץ': 90,
        'ק': 100, 'ר': 200, 'ש': 300, 'ת': 400
    }
    
    total = 0
    for char in hebrew_str:
        if char in values:
            total += values[char]
    
    return total if total > 0 else 999999  # Unknown goes to end


def number_to_hebrew_gematria(num):
    """
    Convert a number to Hebrew gematria notation.
    Examples: 1 → א, 2 → ב, 10 → י, 11 → יא, 20 → כ, 21 → כא, etc.
    """
    if num <= 0:
        return str(num)
    
    # Hebrew letters and their numeric values
    ones = ['', 'א', 'ב', 'ג', 'ד', 'ה', 'ו', 'ז', 'ח', 'ט']  # 0-9
    tens = ['', 'י', 'כ', 'ל', 'מ', 'ן', 'ס', 'ע', 'פ', 'צ']  # 0, 10-90
    hundreds = ['', 'ק', 'ר', 'ש', 'ת']  # 0, 100-400
    
    result = ''
    
    # Handle hundreds
    if num >= 100:
        hundreds_digit = min(num // 100, 4)
        result += hundreds[hundreds_digit]
        num %= 100
    
    # Special cases for 15 and 16 (avoid using God's name)
    if num == 15:
        return result + 'טו'
    elif num == 16:
        return result + 'טז'
    
    # Handle tens
    if num >= 10:
        tens_digit = num // 10
        result += tens[tens_digit]
        num %= 10
    
    # Handle ones
    if num > 0:
        result += ones[num]
    
    return result if result else str(num)


def extract_heading4_info(filename_stem):
    """
    Extract heading 4 information from filename.
    Handles special patterns:
      - "PEREK1" or "perek1" → "פרק א"
      - "PEREK2" → "פרק ב"
      - "PEREK11" → "פרק יא"
      - "PEREK01A" or "perek1a" → "פרק א 1" (letter becomes number, number becomes letter)
      - "MEKOROS" or "MKOROS" → "מקורות"
      - "MEKOROS1" → "מקורות א"
      - "HAKDOMO" or "HAKDOMO1" → "הקדמה" or "הקדמה א"
    Returns the Hebrew string or None if no pattern matched.
    """
    stem = filename_stem.strip().lower()
    
    # Check for MEKOROS/MKOROS with optional number
    mekoros_match = re.match(r'^m?koros0*(\d*)$', stem, re.IGNORECASE)
    if mekoros_match:
        num_str = mekoros_match.group(1)
        if num_str:
            number = int(num_str)
            hebrew_gematria = number_to_hebrew_gematria(number)
            return f"מקורות {hebrew_gematria}"
        else:
            return 'מקורות'
    
    # Check for HAKDOMO with optional number
    hakdomo_match = re.match(r'^hakdomo0*(\d*)$', stem, re.IGNORECASE)
    if hakdomo_match:
        num_str = hakdomo_match.group(1)
        if num_str:
            number = int(num_str)
            hebrew_gematria = number_to_hebrew_gematria(number)
            return f"הקדמה {hebrew_gematria}"
        else:
            return 'הקדמה'
    
    # Pattern: perek followed by number (with optional leading zeros) and optional letter
    perek_match = re.match(r'^perek0*(\d+)([a-z])?$', stem, re.IGNORECASE)
    if perek_match:
        number = int(perek_match.group(1))
        letter = perek_match.group(2)
        
        # Convert number to Hebrew gematria
        hebrew_gematria = number_to_hebrew_gematria(number)
        
        if letter:
            # Convert letter to number (a=1, b=2, etc.)
            letter_num = ord(letter.lower()) - ord('a') + 1
            return f"פרק {hebrew_gematria} {letter_num}"
        else:
            return f"פרק {hebrew_gematria}"
    
    return None


def extract_year(filename_stem):
    """
    Extract year from filename.
    Looks for Hebrew year pattern like תש״כ, תשכ_ז תשכח, תשנ״ט, etc.
    Years always start with תש (taf-shin) and are 3-4 characters long.
    Returns None if no year found (year is optional).
    """
    stem = filename_stem.strip()

    # Split by common separators (including underscore)
    parts = re.split(r"[\s\-–—_]+", stem)
    parts = [p.strip() for p in parts if p.strip()]

    # Look for year pattern: must start with תש and be 3-4 chars total
    # This excludes parshah names like תזריע, תבוא, etc.
    year_pattern = r"^תש[\u0590-\u05FF״׳\"]$|^תש[\u0590-\u05FF״׳\"][\u0590-\u05FF״׳\"]$"

    for part in parts:
        # Check if it matches the תש year pattern and is 3-4 chars
        if re.match(year_pattern, part) and 3 <= len(part) <= 4:
            return part

    # Fallback: look for תש pattern with correct length
    for part in parts:
        if len(part) >= 3 and len(part) <= 4 and part[0:2] == "תש":
            return part

    return None


def extract_year_from_text(text):
    """
    Extract year from text (similar to extract_year but works on paragraph text).
    Looks for Hebrew year pattern like תש״כ, תשכ_ז תשכח, תשנ״ט, etc.
    """
    if not text:
        return None

    # Look for year pattern in the text
    # Pattern: תש followed by 1-2 Hebrew characters or gershayim
    year_pattern = r"תש[\u0590-\u05FF״׳\"][\u0590-\u05FF״׳\"]?"
    matches = re.findall(year_pattern, text)

    for match in matches:
        if 3 <= len(match) <= 4:
            return match

    return None


def is_valid_gematria_number(text):
    """
    Check if a Hebrew text is a valid gematria number (not a regular word).
    Hebrew alphabet numbering: א=1, ב=2, ג=3... or gematria combinations.
    """
    # ALL single Hebrew letters are valid numbers (Hebrew alphabet numbering)
    # א, ב, ג, ד, ה, ו, ז, ח, ט, י, כ, ל, מ, נ, ס, ע, פ, צ, ק, ר, ש, ת
    if len(text) == 1:
        return True
    
    # Exclude common Hebrew WORDS that aren't numbers (multi-letter only)
    non_numbers = {
        "מבוא", "פרק", "חלק", "סימן", "דרוש", "מאמר", "שיחה", 
        "הקדמה", "תוכן", "ענין", "דבר", "מכתב", "נושא", "הערות",
        "הגהות", "ביאור", "פסוק", "דין", "הלכה", "מצוה", "הערה"
    }
    if text in non_numbers:
        return False
    
    # For multi-letter: if it's 2-4 letters and not in blacklist, likely a gematria number
    # Examples: יב (12), כג (23), קנד (154), ריח (218), תשכז (5727)
    return len(text) <= 4


def detect_parshah_boundary(text):
    """
    Detect if a paragraph indicates the start of a new parshah or section.
    Returns (is_boundary, parshah_name, year) tuple.
    """
    if not text:
        return (False, None, None)

    txt = text.strip()

    # Pattern 0: Hebrew letter-number (siman) - like ב, ג, ריב, ריח, etc.
    # These are 1-4 Hebrew letters, possibly followed by period/whitespace
    # representing section numbers
    siman_match = re.match(r"^([א-ת]{1,4})[\.\s\t]*$", txt)
    if siman_match and len(txt) <= 10:
        siman = siman_match.group(1)
        # Validate it's actually a gematria number, not a word
        if is_valid_gematria_number(siman):
            # This is a siman number, use it as the section name (add period for display)
            return (True, f"{siman}.", None)

    # Pattern 1: "פרשת [name]" - explicit parshah marker
    parshah_match = re.match(r"^פרשת\s+([א-ת\s]+?)(?:\s|$)", txt)
    if parshah_match:
        parshah_name = parshah_match.group(1).strip()
        year = extract_year_from_text(txt)
        return (True, parshah_name, year)

    # Pattern 2: "פרשת [name] שנת [year]" or "פרשת [name] - [year]"
    parshah_with_year = re.match(
        r"^פרשת\s+([א-ת\s]+?)(?:\s+שנת|\s*[-–—])\s*(.+?)(?:\s|$)", txt
    )
    if parshah_with_year:
        parshah_name = parshah_with_year.group(1).strip()
        year_text = parshah_with_year.group(2).strip()
        year = extract_year_from_text(year_text) or extract_year_from_text(txt)
        return (True, parshah_name, year)

    # Pattern 3: Just a parshah name (common Hebrew parshah names)
    # This is a fallback - look for known parshah patterns
    # Common parshah names are usually 2-4 Hebrew words
    if re.match(r"^[א-ת\s]{4,30}$", txt) and not re.search(r"[.!?,\[\]\*]", txt):
        # Check if it might be a parshah name (not too long, no punctuation)
        # Extract year if present
        year = extract_year_from_text(txt)
        # If we found a year, this might be a parshah header
        if year:
            # Try to extract parshah name (everything before the year)
            parts = txt.split(year)
            if parts and parts[0].strip():
                parshah_name = parts[0].strip()
                # Remove common prefixes
                parshah_name = re.sub(r"^פרשת\s+", "", parshah_name).strip()
                if parshah_name:
                    return (True, parshah_name, year)

    return (False, None, None)


def convert_multi_parshah_to_json(
    input_path, output_path, book, sefer, skip_parshah_prefix=False
):
    """
    Convert a multi-parshah document to JSON format.
    Each list item becomes a single chunk with a title.
    """
    from datetime import datetime
    
    source = Document(input_path)
    
    current_date = datetime.now().strftime("%Y-%m-%d")
    
    # Create JSON structure
    json_data = {
        "book_name_he": book,
        "book_name_en": "",
        "book_metadata": {"date": current_date, "sefer": sefer},
        "chunks": [],  # Array of chunks, one per list item
    }
    
    print(f"\nScanning {len(source.paragraphs)} paragraphs for List items...")
    
    current_chunk = None
    chunk_id = 0
    list_counter = 0
    current_chunk_paragraphs = []
    
    for i, para in enumerate(source.paragraphs):
        txt = para.text.strip()
        
        # Check if this paragraph is a list item
        is_list_item = False
        section_name = None
        
        try:
            style_name = para.style.name if para.style else None
            if style_name and 'list' in style_name.lower():
                is_list_item = True
                list_counter += 1
                
                # Use the full paragraph text as the chunk title (Heading 3)
                section_name = txt
                if not section_name and i + 1 < len(source.paragraphs):
                    section_name = source.paragraphs[i + 1].text.strip()
                
                if not section_name:
                    section_name = str(list_counter)
                
                print(f"  ✓ Found list item #{list_counter} at paragraph {i}: '{section_name[:60] if len(section_name) > 60 else section_name}'")
        except:
            pass
        
        # If it's a list item, start a new chunk
        if is_list_item:
            # Save previous chunk if exists
            if current_chunk is not None and current_chunk_paragraphs:
                # Combine all paragraphs into the chunk text
                current_chunk["text"] = "\n\n".join(current_chunk_paragraphs)
                json_data["chunks"].append(current_chunk)
            
            # Start new chunk
            chunk_id += 1
            current_chunk = {
                "chunk_id": chunk_id,
                "chunk_metadata": {
                    "chunk_title": section_name  # Heading 3
                },
                "text": ""
            }
            current_chunk_paragraphs = []
        
        # Add non-empty paragraphs to current chunk
        elif txt and current_chunk is not None:
            current_chunk_paragraphs.append(txt)
    
    # Don't forget to add the last chunk
    if current_chunk is not None and current_chunk_paragraphs:
        current_chunk["text"] = "\n\n".join(current_chunk_paragraphs)
        json_data["chunks"].append(current_chunk)
    
    print(f"\nTotal chunks found: {len(json_data['chunks'])}")
    print()
    
    # Write JSON file
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(json_data, f, ensure_ascii=False, indent=2)


def reformat_multi_parshah_docx(
    input_path, output_path, book, sefer, skip_parshah_prefix=False
):
    """
    Process a single document containing multiple parshahs.
    Creates a new document with headings inserted before each list item.
    """
    # Open source document
    source = Document(input_path)
    
    # Create new document
    new_doc = Document()
    configure_styles(new_doc)
    
    print(f"\nScanning {len(source.paragraphs)} paragraphs for List items...")
    
    list_counter = 0
    
    for i, para in enumerate(source.paragraphs):
        txt = para.text.strip()
        
        # Check if this paragraph is a list item
        is_list_item = False
        section_name = None
        
        try:
            style_name = para.style.name if para.style else None
            if style_name and 'list' in style_name.lower():
                is_list_item = True
                list_counter += 1
                
                # Use the full paragraph text as the section name
                # Check this paragraph first, if empty check next paragraph
                section_name = txt
                if not section_name and i + 1 < len(source.paragraphs):
                    section_name = source.paragraphs[i + 1].text.strip()
                
                if not section_name:
                    # Fallback: use counter
                    section_name = str(list_counter)
                
                print(f"  ✓ Found list item #{list_counter} at paragraph {i}: '{section_name[:60] if len(section_name) > 60 else section_name}'")
        except:
            pass
        
        # If it's a list item, insert headings first
        if is_list_item:
            # Prepare heading texts
            parshah_heading = (
                section_name if skip_parshah_prefix else f"פרשת {section_name}"
            )
            
            headings = [
                ("Heading 1", book),
                ("Heading 2", sefer),
                ("Heading 3", parshah_heading),
            ]
            
            # Insert the headings
            for level, text in headings:
                if text:
                    p = new_doc.add_paragraph(text, style=level)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.right_to_left = True
        
        # Now copy the original paragraph
        new_p = new_doc.add_paragraph()
        
        # Copy paragraph-level formatting
        pf_source = para.paragraph_format
        pf_new = new_p.paragraph_format
        
        # Copy alignment
        if para.alignment:
            new_p.alignment = para.alignment
        else:
            new_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Set RTL
        try:
            pf_new.right_to_left = True
        except:
            pass
        
        # Copy indentation
        if pf_source.left_indent is not None:
            pf_new.left_indent = pf_source.left_indent
        if pf_source.right_indent is not None:
            pf_new.right_indent = pf_source.right_indent
        if pf_source.first_line_indent is not None:
            pf_new.first_line_indent = pf_source.first_line_indent
        
        # Copy spacing
        pf_new.space_before = pf_source.space_before
        pf_new.space_after = pf_source.space_after
        pf_new.line_spacing = pf_source.line_spacing
        if pf_source.line_spacing_rule is not None:
            pf_new.line_spacing_rule = pf_source.line_spacing_rule
        
        # Copy all runs
        for run in para.runs:
            new_r = new_p.add_run(run.text)
            
            # Copy font properties
            if run.font.bold is not None:
                new_r.font.bold = run.font.bold
            if run.font.italic is not None:
                new_r.font.italic = run.font.italic
            if run.font.underline is not None:
                new_r.font.underline = run.font.underline
            if run.font.size is not None:
                new_r.font.size = run.font.size
            if run.font.name is not None:
                new_r.font.name = run.font.name
            if run.font.color.rgb is not None:
                new_r.font.color.rgb = run.font.color.rgb
    
    print(f"\nTotal list items found: {list_counter}")
    print()
    
    new_doc.save(output_path)


def reformat_docx(
    input_path, output_path, book, sefer, parshah, filename, skip_parshah_prefix=False
):
    source = Document(input_path)
    new_doc = Document()
    configure_styles(new_doc)

    # Add document headings
    parshah_heading = parshah if skip_parshah_prefix else f"פרשת {parshah}"

    # Heading 4 is optional - only add if filename is provided
    headings = [
        ("Heading 1", book),
        ("Heading 2", sefer),
        ("Heading 3", parshah_heading),
    ]
    
    if filename:
        headings.append(("Heading 4", filename))

    for level, text in headings:
        if not text:
            continue
        p = new_doc.add_paragraph(text, style=level)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Ensure RTL for Hebrew text
        p.paragraph_format.right_to_left = True

    # Process body text with smart header skipping
    in_header_section = True

    for para in source.paragraphs:
        # Get the full paragraph text including ALL characters
        full_text = para.text
        txt = full_text.strip()

        # If we're still in the header section
        if in_header_section:
            # Check if this looks like substantial content
            if txt and should_start_content(txt):
                in_header_section = False
                # Fall through to copy this paragraph
            # Skip if it's an old header
            elif txt and is_old_header(txt):
                continue
            # Skip empty paragraphs in header section
            elif not txt:
                continue
            else:
                # Non-header, non-empty text in header section - shouldn't happen but be safe
                continue

        # After header section started
        # Skip only matching old headers, preserve EVERYTHING else including empty paragraphs
        if txt and is_old_header(txt):
            continue

        # Copy the entire paragraph element to preserve ALL formatting
        # This includes empty paragraphs which create spacing
        new_p = new_doc.add_paragraph()

        # Copy paragraph-level formatting
        pf_source = para.paragraph_format
        pf_new = new_p.paragraph_format

        # Check if this is a separator (only asterisks)
        cleaned_text = txt.replace(' ', '').replace('\u00a0', '')
        is_separator = cleaned_text and all(c == '*' for c in cleaned_text)
        
        # Center separators, preserve centered alignment, otherwise force RTL right alignment
        if is_separator or para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            new_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Set RTL direction for Hebrew
        pf_new.right_to_left = True

        # Copy all paragraph format attributes
        if pf_source.left_indent is not None:
            pf_new.left_indent = pf_source.left_indent
        if pf_source.right_indent is not None:
            pf_new.right_indent = pf_source.right_indent
        if pf_source.first_line_indent is not None:
            pf_new.first_line_indent = pf_source.first_line_indent

        # Always copy spacing - these are critical for layout
        pf_new.space_before = pf_source.space_before
        pf_new.space_after = pf_source.space_after
        pf_new.line_spacing = pf_source.line_spacing
        if pf_source.line_spacing_rule is not None:
            pf_new.line_spacing_rule = pf_source.line_spacing_rule

        # Copy keep together settings
        if pf_source.keep_together is not None:
            pf_new.keep_together = pf_source.keep_together
        if pf_source.keep_with_next is not None:
            pf_new.keep_with_next = pf_source.keep_with_next
        if pf_source.page_break_before is not None:
            pf_new.page_break_before = pf_source.page_break_before
        if pf_source.widow_control is not None:
            pf_new.widow_control = pf_source.widow_control

        # Copy ALL runs including ones with just symbols/whitespace
        for run in para.runs:
            # Copy the run even if it's just whitespace or symbols
            new_r = new_p.add_run(run.text)

            # Copy all font properties
            if run.font.bold is not None:
                new_r.font.bold = run.font.bold
            if run.font.italic is not None:
                new_r.font.italic = run.font.italic
            if run.font.underline is not None:
                new_r.font.underline = run.font.underline
            if run.font.size is not None:
                new_r.font.size = run.font.size
            if run.font.name is not None:
                new_r.font.name = run.font.name
            if run.font.color.rgb is not None:
                new_r.font.color.rgb = run.font.color.rgb
            if run.font.highlight_color is not None:
                new_r.font.highlight_color = run.font.highlight_color
            if run.font.all_caps is not None:
                new_r.font.all_caps = run.font.all_caps
            if run.font.small_caps is not None:
                new_r.font.small_caps = run.font.small_caps
            if run.font.strike is not None:
                new_r.font.strike = run.font.strike
            if run.font.superscript is not None:
                new_r.font.superscript = run.font.superscript
            if run.font.subscript is not None:
                new_r.font.subscript = run.font.subscript

        # Add a blank line after each paragraph with content
        if txt:
            new_doc.add_paragraph()

    new_doc.save(output_path)


def convert_to_json(
    input_path, output_path, book, sefer, title, filename, skip_parshah_prefix=False
):
    """
    Convert docx to JSON structure with chunks.
    Each paragraph becomes a chunk.
    """
    source = Document(input_path)

    # Get current date
    from datetime import datetime

    current_date = datetime.now().strftime("%Y-%m-%d")

    # Create JSON structure
    # Use filename (heading4) for book_name_he, fallback to title if not provided
    json_data = {
        "book_name_he": filename if filename else title,
        "book_name_en": "",
        "book_metadata": {"date": current_date},
        "chunks": [],
    }

    # Process body text with smart header skipping
    in_header_section = True
    chunk_id = 1

    for para in source.paragraphs:
        full_text = para.text
        txt = full_text.strip()

        # If we're still in the header section
        if in_header_section:
            # Check if this looks like substantial content
            if txt and should_start_content(txt):
                in_header_section = False
                # Fall through to include this paragraph
            # Skip if it's an old header
            elif txt and is_old_header(txt):
                continue
            # Skip empty paragraphs in header section
            elif not txt:
                continue
            else:
                continue

        # After header section, skip old headers but keep everything else
        if txt and is_old_header(txt):
            continue

        # Only add non-empty paragraphs as chunks
        if txt:
            chunk = {"chunk_id": chunk_id, "chunk_metadata": {
                "chunk_title": f"Paragraph {chunk_id}"
            }, "text": txt}
            json_data["chunks"].append(chunk)
            chunk_id += 1

    # Write JSON file
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(json_data, f, ensure_ascii=False, indent=2)


# -------------------------------
# Main CLI entry point
# -------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Reformat Hebrew DOCX files to standardized schema."
    )
    parser.add_argument("--book", help="Book title (Heading 1). Optional in DAF mode when using folder structure.")
    parser.add_argument(
        "--sefer",
        help="Sefer/tractate title (Heading 2). If not provided, uses folder name.",
    )
    parser.add_argument(
        "--parshah",
        help="Parshah name (Heading 3). If not provided, uses subfolder names.",
    )
    parser.add_argument(
        "--skip-parshah-prefix",
        action="store_true",
        help="Skip adding 'פרשת' prefix to parshah name in Heading 3",
    )
    parser.add_argument(
        "--json",
        action="store_true",
        help="Output as JSON structure instead of formatted Word documents",
    )
    parser.add_argument(
        "--docs",
        default="docs",
        help="Input folder containing .docx files or subfolders (or single file for multi-parshah mode)",
    )
    parser.add_argument("--out", default="output", help="Output folder")
    parser.add_argument(
        "--multi-parshah",
        action="store_true",
        help="Process a single document containing multiple parshahs. Detects parshah boundaries and inserts headings.",
    )
    parser.add_argument(
        "--combine-parshah",
        action="store_true",
        help="Combine all year documents per parshah into one Word file with four headings per year.",
    )
    parser.add_argument(
        "--daf-mode",
        action="store_true",
        help="DAF mode: For IDML files, use page markers as Heading 3 (דף) and Heading 4 (עמוד).",
    )

    # --- Top-level: combine all year docs per parshah into one doc ---
    def combine_parshah_docs(subdir, out_subdir, book, sefer, parshah, skip_parshah_prefix, daf_mode=False):
        files = get_processable_files(subdir)
        if not files:
            return
        from docx import Document
        combined_doc = Document()
        configure_styles(combined_doc)
        
        # Track previous heading values to only insert when they change
        prev_book = None
        prev_sefer = None
        prev_parshah = None
        prev_heading4 = None
        
        for path in sorted(files):
            temp_docx = None
            needs_cleanup = False
            try:
                filename_stem = Path(path).stem if path.suffix else path.name
                title = filename_stem.replace('-formatted', '')
                
                # Try to extract year, then heading4 info, then use title
                year = extract_year(filename_stem)
                heading4_info = extract_heading4_info(filename_stem)
                heading4 = year or heading4_info or title
                
                # Prepare heading values
                current_parshah = parshah if skip_parshah_prefix else f"פרשת {parshah}"
                
                # Convert to .docx if needed
                input_path, needs_cleanup = convert_to_docx(path, daf_mode=daf_mode)
                if needs_cleanup:
                    temp_docx = input_path
                
                source = Document(input_path)
                
                # In DAF mode, add Book/Sefer as Heading 1/2, then copy דף/עמוד from content
                if daf_mode:
                    # Add Book and Sefer headings only when they change (or first file)
                    if prev_book != book:
                        p = combined_doc.add_paragraph(book, style="Heading 1")
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        p.paragraph_format.right_to_left = True
                        prev_book = book
                    
                    if prev_sefer != sefer:
                        p = combined_doc.add_paragraph(sefer, style="Heading 2")
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        p.paragraph_format.right_to_left = True
                        prev_sefer = sefer
                    
                    # Copy all paragraphs including Heading 3 (דף) and Heading 4 (עמוד)
                    for para in source.paragraphs:
                        txt = para.text.strip()
                        if not txt:  # Skip empty paragraphs
                            continue
                        
                        # Check if this is a separator (stars only)
                        cleaned_text = txt.replace(' ', '').replace('\u00a0', '')
                        is_separator = cleaned_text and all(c == '*' for c in cleaned_text)
                        
                        # Copy paragraph with its style
                        if para.style.name.startswith('Heading'):
                            new_p = combined_doc.add_paragraph(para.text, style=para.style.name)
                            new_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            new_p.paragraph_format.right_to_left = True
                        elif is_separator:
                            # Separator should be centered
                            new_p = combined_doc.add_paragraph(para.text)
                            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            new_p = combined_doc.add_paragraph()
                            for run in para.runs:
                                new_r = new_p.add_run(run.text)
                                if run.font.bold is not None:
                                    new_r.font.bold = run.font.bold
                                if run.font.italic is not None:
                                    new_r.font.italic = run.font.italic
                                if run.font.underline is not None:
                                    new_r.font.underline = run.font.underline
                            new_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            try:
                                new_p.paragraph_format.right_to_left = True
                            except:
                                pass
                else:
                    # Non-DAF mode: Add standard headings only when they change
                    if prev_book != book:
                        p = combined_doc.add_paragraph(book, style="Heading 1")
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        p.paragraph_format.right_to_left = True
                        prev_book = book
                    
                    if prev_sefer != sefer:
                        p = combined_doc.add_paragraph(sefer, style="Heading 2")
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        p.paragraph_format.right_to_left = True
                        prev_sefer = sefer
                    
                    if prev_parshah != current_parshah:
                        p = combined_doc.add_paragraph(current_parshah, style="Heading 3")
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        p.paragraph_format.right_to_left = True
                        prev_parshah = current_parshah
                    
                    # Add Heading 4 only when it changes (or first file)
                    if heading4 and prev_heading4 != heading4:
                        p = combined_doc.add_paragraph(heading4, style="Heading 4")
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        p.paragraph_format.right_to_left = True
                        prev_heading4 = heading4
                    
                    # Copy paragraphs from source, skipping old headers
                    in_header_section = True
                    for para in source.paragraphs:
                        full_text = para.text
                        txt = full_text.strip()
                        # If we're still in the header section
                        if in_header_section:
                            if txt and not is_old_header(txt):
                                in_header_section = False
                            elif txt and is_old_header(txt):
                                continue
                            elif not txt:
                                continue
                            else:
                                continue
                        # After header section started
                        if txt and is_old_header(txt):
                            continue
                        # Copy paragraph
                        new_p = combined_doc.add_paragraph()
                        pf_source = para.paragraph_format
                        pf_new = new_p.paragraph_format
                        if para.alignment:
                            new_p.alignment = para.alignment
                        else:
                            new_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        try:
                            pf_new.right_to_left = True
                        except:
                            pass
                        if pf_source.left_indent is not None:
                            pf_new.left_indent = pf_source.left_indent
                        if pf_source.right_indent is not None:
                            pf_new.right_indent = pf_source.right_indent
                        if pf_source.first_line_indent is not None:
                            pf_new.first_line_indent = pf_source.first_line_indent
                        pf_new.space_before = pf_source.space_before
                        pf_new.space_after = pf_source.space_after
                        pf_new.line_spacing = pf_source.line_spacing
                        if pf_source.line_spacing_rule is not None:
                            pf_new.line_spacing_rule = pf_source.line_spacing_rule
                        for run in para.runs:
                            new_r = new_p.add_run(run.text)
                            if run.font.bold is not None:
                                new_r.font.bold = run.font.bold
                            if run.font.italic is not None:
                                new_r.font.italic = run.font.italic
                            if run.font.underline is not None:
                                new_r.font.underline = run.font.underline
                            if run.font.size is not None:
                                new_r.font.size = run.font.size
                            if run.font.name is not None:
                                new_r.font.name = run.font.name
                            if run.font.color.rgb is not None:
                                new_r.font.color.rgb = run.font.color.rgb
                
                # Add a blank line after each file
                combined_doc.add_paragraph()
            finally:
                if temp_docx and temp_docx.exists():
                    temp_docx.unlink()
        # Save combined document
        # In DAF mode, use sefer name for combined file; in normal mode use parshah
        if daf_mode:
            out_name = f"{sefer}-combined.docx"
        else:
            out_name = f"{parshah}-combined.docx"
        out_path = out_subdir / out_name
        combined_doc.save(out_path)
        print(f"  ✓ Combined {len(files)} file(s) into {out_path}")
    args = parser.parse_args()

    docs_path = Path(args.docs)
    out_dir = Path(args.out)

    # Validate --book requirement
    # In DAF mode with folder structure (no --sefer/--parshah), --book is optional (derived from folder)
    # In all other cases, --book is required
    if not args.book:
        if not args.daf_mode or args.sefer or args.parshah:
            print("Error: --book is required (except in DAF mode with folder structure)")
            return

    # Multi-parshah mode: process a single file
    if args.multi_parshah:
        if not docs_path.exists():
            print(f"Error: Input file '{docs_path}' does not exist")
            return

        if docs_path.is_dir():
            print(
                f"Error: '{docs_path}' is a directory. For multi-parshah mode, provide a single file."
            )
            return

        # Check if sefer is provided (required for multi-parshah mode)
        if not args.sefer:
            print("Error: --sefer is required for multi-parshah mode")
            return

        out_dir.mkdir(exist_ok=True)

        # Determine output filename
        input_stem = docs_path.stem
        if args.json:
            out_name = f"{input_stem}.json"
            out_path = out_dir / out_name
        else:
            out_name = f"{input_stem.replace('-formatted', '')}-formatted.docx"
            out_path = out_dir / out_name

        file_display_name = docs_path.name
        print(f"📚 Processing multi-parshah document: {file_display_name}\n")
        print(f"   Book: {args.book}")
        print(f"   Sefer: {args.sefer}\n")

        temp_docx = None
        needs_cleanup = False
        try:
            # Convert to .docx format
            input_path, needs_cleanup = convert_to_docx(docs_path, daf_mode=args.daf_mode)
            if needs_cleanup:
                temp_docx = input_path
                print("Converting to .docx format... done\n")

            print("Processing... ", end="")
            if args.json:
                convert_multi_parshah_to_json(
                    input_path, out_path, args.book, args.sefer, args.skip_parshah_prefix
                )
            else:
                reformat_multi_parshah_docx(
                    input_path, out_path, args.book, args.sefer, args.skip_parshah_prefix
                )
            print("✓ done")
            print(f"\n✅ Output saved to: {out_path}")
        except Exception as e:
            print(f"⚠️ error: {e}")
            traceback.print_exc()
        finally:
            if temp_docx and temp_docx.exists():
                temp_docx.unlink()
        return

    # Regular mode: process directory
    docs_dir = docs_path

    # Check if docs_dir exists
    if not docs_dir.exists():
        print(f"Error: Input directory '{docs_dir}' does not exist")
        return

    if not docs_dir.is_dir():
        print(f"Error: '{docs_dir}' is not a directory")
        return

    out_dir.mkdir(exist_ok=True)

    # Create json subdirectory if needed
    if args.json:
        (out_dir / "json").mkdir(exist_ok=True)

    # Check if using folder structure mode (no sefer/parshah specified)
    if not args.sefer and not args.parshah:
        # In DAF mode: docs_dir is the collection (Heading 1), subdirs are masechet (Heading 2)
        # In normal mode: docs_dir is sefer, subdirs are parshah
        
        if args.daf_mode:
            # DAF mode: parent folder = Heading 1 (book), current folder = Heading 2 (masechet/sefer)
            book = docs_dir.name  # e.g., "אגדות מהריט"
            print(f"📚 Processing DAF mode structure: {book}\n")
        else:
            # Normal mode: folder = sefer, use --book argument
            book = args.book
            sefer = docs_dir.name
            print(f"📚 Processing folder structure: {sefer}\n")

        # Get all subdirectories
        subdirs = [d for d in docs_dir.iterdir() if d.is_dir()]

        if not subdirs:
            print(f"No subdirectories found in {docs_dir}")
            return

        total_success = 0
        total_files = 0

        for subdir in subdirs:
            if args.daf_mode:
                # In DAF mode: subdir is the masechet (Heading 2)
                sefer = subdir.name  # e.g., "ביצה"
                parshah = None  # Not used in DAF mode
            else:
                # In normal mode: subdir is the parshah (Heading 3)
                parshah = subdir.name
            
            # Create output subdirectory
            if args.json:
                out_subdir = out_dir / "json" / book / sefer if args.daf_mode else out_dir / "json" / sefer / parshah
            else:
                out_subdir = out_dir / book / sefer if args.daf_mode else out_dir / sefer / parshah
            out_subdir.mkdir(parents=True, exist_ok=True)

            if args.combine_parshah:
                display_name = sefer if args.daf_mode else parshah
                print(f"📂 Combining {display_name} ...")
                combine_parshah_docs(subdir, out_subdir, book, sefer, parshah, args.skip_parshah_prefix, daf_mode=args.daf_mode)
                total_success += 1
                continue

            files = get_processable_files(subdir)
            if not files:
                continue

            display_name = sefer if args.daf_mode else parshah
            print(f"📂 {display_name} ({len(files)} file(s))")

            for i, path in enumerate(files, 1):
                temp_docx = None
                needs_cleanup = False
                try:
                    filename_stem = Path(path).stem if path.suffix else path.name
                    title = filename_stem.replace("-formatted", "")
                    
                    # Try to extract year, then heading4 info, then use title
                    year = extract_year(title)
                    heading4_info = extract_heading4_info(title)
                    
                    # Determine heading 4 text (optional)
                    heading4 = year or heading4_info or title
                    
                    if args.json:
                        out_name = f"{filename_stem}.json"
                        out_path = out_subdir / out_name
                    else:
                        out_name = f"{filename_stem.replace('-formatted', '')}-formatted.docx"
                        out_path = out_subdir / out_name
                    
                    file_display_name = path.stem if path.suffix else path.name
                    print(f"  [{i}/{len(files)}] {file_display_name} → {out_path.name} ...", end=" ")
                    
                    # Convert to .docx format
                    input_path, needs_cleanup = convert_to_docx(path, daf_mode=args.daf_mode)
                    if needs_cleanup:
                        temp_docx = input_path
                    
                    # In DAF mode with IDML files, skip reformatting (headings already correct)
                    if args.daf_mode and path.suffix.lower() == '.idml':
                        from docx import Document
                        doc = Document(input_path)
                        doc.save(out_path)
                    elif args.json:
                        convert_to_json(
                            input_path,
                            out_path,
                            book if args.daf_mode else args.book,
                            sefer,
                            title,
                            heading4,
                            args.skip_parshah_prefix,
                        )
                    else:
                        reformat_docx(
                            input_path,
                            out_path,
                            book if args.daf_mode else args.book,
                            sefer,
                            parshah,
                            heading4,
                            args.skip_parshah_prefix,
                        )
                    print("✓ done")
                    total_success += 1
                    total_files += 1
                except Exception as e:
                    print(f"⚠️ error: {e}")
                    total_files += 1
                finally:
                    if temp_docx and temp_docx.exists():
                        temp_docx.unlink()
            print()
        print(f"✅ All done. Successfully processed {total_success}/{total_files} file(s).")
        return

    # Original single folder mode
    if not args.sefer or not args.parshah:
        print(
            "Error: Both --sefer and --parshah are required when not using folder structure mode"
        )
        return

    # Collect processable files
    files = get_processable_files(docs_dir)
    if not files:
        print(f"No supported files (.doc, .docx, .idml, or DOS-encoded) found in {docs_dir}")
        return

    print(f"📚 Processing {len(files)} file(s)...\n")

    success_count = 0
    for i, path in enumerate(files, 1):
        temp_docx = None
        needs_cleanup = False
        try:
            # Extract information from ORIGINAL filename (before any conversion)
            filename_stem = Path(path).stem if path.suffix else path.name
            title = filename_stem.replace("-formatted", "")
            
            # Try to extract year, then heading4 info, then use title
            year = extract_year(filename_stem)
            heading4_info = extract_heading4_info(filename_stem)
            heading4 = year or heading4_info or title

            # Output filename based on format
            if args.json:
                out_name = f"{filename_stem}.json"
                out_path = out_dir / "json" / out_name
            else:
                out_name = f"{filename_stem.replace('-formatted', '')}-formatted.docx"
                out_path = out_dir / out_name

            file_display_name = path.stem if path.suffix else path.name
            print(
                f"[{i}/{len(files)}] Processing {file_display_name} → {out_path.name} ...",
                end=" ",
            )

            # Convert to .docx format
            input_path, needs_cleanup = convert_to_docx(path, daf_mode=args.daf_mode)
            if needs_cleanup:
                temp_docx = input_path

            # In DAF mode with IDML files, the conversion already added the correct headings
            # So we just copy the converted file directly without reformatting
            if args.daf_mode and path.suffix.lower() == '.idml':
                from docx import Document
                doc = Document(input_path)
                doc.save(out_path)
            # Pass the heading4 info
            elif args.json:
                convert_to_json(
                    input_path,
                    out_path,
                    args.book,
                    args.sefer,
                    title,
                    heading4,
                    args.skip_parshah_prefix,
                )
            else:
                reformat_docx(
                    input_path,
                    out_path,
                    args.book,
                    args.sefer,
                    args.parshah,
                    heading4,
                    args.skip_parshah_prefix,
                )
            print("✓ done")
            success_count += 1

        except Exception as e:
            print(f"⚠️ error: {e}")
        finally:
            # Clean up temp file
            if temp_docx and temp_docx.exists():
                temp_docx.unlink()

    print(
        f"\n✅ All done. Successfully processed {success_count}/{len(files)} file(s)."
    )


if __name__ == "__main__":
    main()
